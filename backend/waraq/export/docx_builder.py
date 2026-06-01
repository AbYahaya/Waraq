"""T-9.2.1 — Translation export DOCX builder.

Per Formatvorlagen-Baseline v1.1 §7.2:
- Per-paragraph RTL marking on Arabic text (Sprint 5 RTL-Per-Run-Test).
- Heading styles per `\\o "1-6"` TOC depth (Schluss-Audit Paket 7 Item 2 (a)).
- Footnotes per `eachSect`.
- Page setup: A4, generous margins, header/footer with project title.
- Word-compatible — opens without warnings or repair indicators
  (Word-Kompatibel-Oeffnungs-Test).

Read-only: this module never writes Revision rows, never modifies
Segment text, never writes TRANSLATION-PO rows. Pure read of segment
target text + minimal page/block metadata.

The artefact is built fully in memory; bytes are returned to the
caller. Persistence to the artefact store happens in the atomic-commit
step (a) inside `run_export_job`.
"""

from __future__ import annotations

import hashlib
import io
import re
import uuid as _uuid
from collections.abc import Mapping, Sequence
from dataclasses import dataclass, field
from typing import Any, Literal, cast

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from waraq.identity.service import new_uuid
from waraq.schemas import Block, Page, Revision, Segment
from waraq.style_profile import DEFAULT_STYLE_PROFILE, normalize_style_profile
from waraq.text_state import (
    resolve_segment_source_text,
    resolve_segment_text_state,
    split_source_target_text,
)
from waraq.translation_styles import read_translation_style_map

TocPosition = Literal["front", "back"]
InlineAlignment = Literal["left", "center", "justify"]


@dataclass(frozen=True, kw_only=True, slots=True)
class TranslationDocxConfig:
    """Export-layout choices captured during preflight."""

    header_heading_level: int = 1
    chapter_break_heading_level: int = 1
    toc_position: TocPosition = "front"
    # Legacy field name kept for preflight/export protocol compatibility.
    # It now controls whether Arabic/OCR source text is included at all.
    display_arabic_chapter_headings: bool = True
    style_profile: Mapping[str, Any] = field(default_factory=dict)


DEFAULT_DOCX_CONFIG = TranslationDocxConfig()

ARABIC_FONT = str(DEFAULT_STYLE_PROFILE["docx_arabic_font_family"])


@dataclass(frozen=True, kw_only=True, slots=True)
class TranslationDocxArtefact:
    """The DOCX bytes plus identity + protocol data."""

    artefact_uuid: _uuid.UUID
    bytes_: bytes
    sha256: str
    size_bytes: int
    n_pages_exported: int
    n_segments_exported: int
    block_types_present: list[str]
    exported_segment_uuids: list[_uuid.UUID]
    exported_page_uuids: list[_uuid.UUID]
    exported_block_uuids: list[_uuid.UUID]
    protocol: dict[str, Any] = field(default_factory=dict)


def _set_paragraph_rtl(paragraph: Any) -> None:
    """Mark a paragraph right-to-left at the paragraph-properties level
    so each Arabic paragraph carries its own `<w:bidi/>` per
    Formatvorlagen-Baseline v1.1 §7.2 — not relying on document-global
    setting. Sprint 5 RTL-Per-Run-Test asserts this structurally."""
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "right")

    text_direction = p_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        p_pr.append(text_direction)
    text_direction.set(qn("w:val"), "rlTb")

    rtl = p_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        p_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def _set_document_rtl(document: Any) -> None:
    settings = document.settings
    element = settings.element
    rtl = element.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        element.append(rtl)
    rtl.set(qn("w:val"), "1")

    lang = element.find(qn("w:themeFontLang"))
    if lang is None:
        lang = OxmlElement("w:themeFontLang")
        element.append(lang)
    lang.set(qn("w:val"), "ar-SA")
    lang.set(qn("w:bidi"), "ar-SA")


def _ensure_arabic_style(document: Any, base_style_name: str) -> str:
    name = f"Waraq Arabic {base_style_name}"
    try:
        document.styles[name]
        return name
    except KeyError:
        pass

    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    try:
        style.base_style = document.styles[base_style_name]
    except KeyError:
        style.base_style = document.styles["Normal"]
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p_pr = style._element.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")

    text_direction = p_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        p_pr.append(text_direction)
    text_direction.set(qn("w:val"), "rlTb")

    rtl = p_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        p_pr.append(rtl)
    rtl.set(qn("w:val"), "1")

    return name


def _set_run_rtl(
    run: Any, *, font_name: str = ARABIC_FONT, font_size_pt: int | None = None
) -> None:
    """Mark a run right-to-left. Some Word renderers ignore paragraph-level
    `<w:bidi/>` unless runs are also RTL."""
    r_pr = run._r.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")

    complex_script = r_pr.find(qn("w:cs"))
    if complex_script is None:
        complex_script = OxmlElement("w:cs")
        r_pr.append(complex_script)

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:hint"), "cs")

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ar-SA")
    lang.set(qn("w:bidi"), "ar-SA")

    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def _effective_style_profile(config: TranslationDocxConfig) -> dict[str, Any]:
    return normalize_style_profile(dict(config.style_profile or {}))


def _block_style_kind(block_type: str | None) -> str:
    normalized = (block_type or "").strip().lower()
    if normalized in {"ue", "hd", "heading"}:
        return "heading"
    if normalized in {"fn", "footnote"}:
        return "footnote"
    if normalized in {"quran", "hadith"}:
        return "protected"
    if normalized in {"qr", "quote", "marginalia", "rn", "caption"}:
        return "quote"
    return "body"


def _docx_font_size_for_block(
    profile: Mapping[str, Any],
    *,
    block_type: str | None,
    source: bool,
) -> int:
    kind = _block_style_kind(block_type)
    if kind == "heading":
        return int(profile["docx_heading_font_size_pt"])
    if kind == "quote":
        return int(profile["docx_quote_font_size_pt"])
    if kind == "footnote":
        return int(profile["docx_footnote_font_size_pt"])
    if kind == "protected":
        return int(profile["docx_protected_font_size_pt"])
    key = "docx_arabic_font_size_pt" if source else "docx_translation_font_size_pt"
    return int(profile[key])


def _docx_spacing_for_block(profile: Mapping[str, Any], block_type: str | None) -> int:
    kind = _block_style_kind(block_type)
    if kind == "heading":
        return max(int(profile["docx_paragraph_spacing_pt"]), 10)
    if kind == "footnote":
        return min(int(profile["docx_paragraph_spacing_pt"]), 4)
    return int(profile["docx_paragraph_spacing_pt"])


def _apply_block_indent(paragraph: Any, block_type: str | None) -> None:
    kind = _block_style_kind(block_type)
    if kind in {"quote", "protected"}:
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.right_indent = Cm(0.6)


def _apply_arabic_layout(
    paragraph: Any, config: TranslationDocxConfig, *, block_type: str | None = None
) -> None:
    profile = _effective_style_profile(config)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.space_after = Pt(_docx_spacing_for_block(profile, block_type))
    paragraph.paragraph_format.line_spacing = float(profile["docx_line_spacing"])
    _apply_block_indent(paragraph, block_type)
    _set_paragraph_rtl(paragraph)
    for run in paragraph.runs:
        _set_run_rtl(
            run,
            font_name=str(profile["docx_arabic_font_family"]),
            font_size_pt=_docx_font_size_for_block(profile, block_type=block_type, source=True),
        )


def _apply_translation_layout(
    paragraph: Any,
    config: TranslationDocxConfig,
    *,
    block_type: str | None = None,
    alignment: InlineAlignment | None = None,
    internal_style_key: str | None = None,
) -> None:
    profile = _effective_style_profile(config)
    template = _style_template(profile, internal_style_key)
    paragraph.alignment = (
        _docx_alignment(alignment)
        if alignment
        else _docx_alignment_from_string(template.get("alignment") if template else None)
    )
    paragraph.paragraph_format.space_after = Pt(
        float(template.get("paragraph_spacing_px")) * 0.75
        if template and template.get("paragraph_spacing_px") is not None
        else _docx_spacing_for_block(profile, block_type)
    )
    paragraph.paragraph_format.line_spacing = float(
        template.get("line_height") if template else profile["docx_line_spacing"]
    )
    if template:
        paragraph.paragraph_format.first_line_indent = Pt(
            float(template.get("first_line_indent_px", 0)) * 0.75
        )
        paragraph.paragraph_format.left_indent = Pt(
            float(template.get("left_indent_px", 0)) * 0.75
        )
    else:
        _apply_block_indent(paragraph, block_type)
    for run in paragraph.runs:
        run.font.name = str(
            template.get("font_family") if template else profile["docx_translation_font_family"]
        )
        run.font.size = Pt(
            int(
                template.get("docx_font_size_pt")
                if template
                else _docx_font_size_for_block(profile, block_type=block_type, source=False)
            )
        )
        if template:
            run.bold = bool(template.get("bold", False))
            run.italic = bool(template.get("italic", False))


_ALIGNMENT_MARKER_RE = re.compile(
    r"\[\[(left|center|justify)\]\](.*?)\[\[/\1\]\]",
    flags=re.DOTALL,
)
_PARAGRAPH_STYLE_MARKER_RE = re.compile(r"^\[\[style:([a-z0-9_]+)\]\]\s*")
_INLINE_MARKER_TOKEN_RE = re.compile(
    r"\[\[(font:[^\]]+|/font|size:\d+(?:\.\d+)?|/size|/?(?:b|i|u|s))\]\]"
)


def _docx_alignment(alignment: InlineAlignment | None) -> WD_PARAGRAPH_ALIGNMENT:
    if alignment == "center":
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    if alignment == "justify":
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return WD_PARAGRAPH_ALIGNMENT.LEFT


def _docx_alignment_from_string(value: object) -> WD_PARAGRAPH_ALIGNMENT:
    if value == "center":
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    if value == "right":
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    if value == "justify":
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return WD_PARAGRAPH_ALIGNMENT.LEFT


def _style_template(profile: Mapping[str, Any], style_key: str | None) -> Mapping[str, Any] | None:
    if not style_key:
        return None
    templates = profile.get("translation_style_templates")
    if not isinstance(templates, Mapping):
        return None
    template = templates.get(style_key)
    return template if isinstance(template, Mapping) else None


@dataclass(frozen=True, slots=True)
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    font_family: str | None = None
    font_size_px: float | None = None


@dataclass(frozen=True, slots=True)
class TranslationParagraphBlock:
    text: str
    alignment: InlineAlignment | None
    style_key: str | None


def _iter_aligned_blocks(
    text: str, *, style_key: str | None = None
) -> list[TranslationParagraphBlock]:
    blocks: list[tuple[str, InlineAlignment | None]] = []
    last = 0
    for match in _ALIGNMENT_MARKER_RE.finditer(text):
        if match.start() > last:
            blocks.append((text[last : match.start()], None))
        blocks.append((match.group(2), cast(InlineAlignment, match.group(1))))
        last = match.end()
    if last < len(text):
        blocks.append((text[last:], None))
    return [
        TranslationParagraphBlock(block, alignment, style_key)
        for block, alignment in blocks
        if block.strip()
    ]


def _iter_styled_translation_blocks(
    text: str, *, fallback_style_key: str | None = None
) -> list[TranslationParagraphBlock]:
    normalized = text.replace("\r\n", "\n")
    parts = [part.strip() for part in re.split(r"\n+", normalized) if part.strip()]
    if not parts and text.strip():
        parts = [text.strip()]
    blocks: list[TranslationParagraphBlock] = []
    for raw in parts:
        match = _PARAGRAPH_STYLE_MARKER_RE.match(raw)
        style_key = match.group(1) if match else fallback_style_key
        content = raw[match.end() :] if match else raw
        blocks.extend(_iter_aligned_blocks(content, style_key=style_key))
    return blocks


def _iter_inline_runs(text: str) -> list[InlineRun]:
    runs: list[InlineRun] = []
    bold = italic = underline = strike = False
    font_family: str | None = None
    font_size_px: float | None = None
    cursor = 0
    for match in _INLINE_MARKER_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            runs.append(
                InlineRun(
                    text[cursor : match.start()],
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strike=strike,
                    font_family=font_family,
                    font_size_px=font_size_px,
                )
            )
        token = match.group(1)
        if token == "b":
            bold = True
        elif token == "/b":
            bold = False
        elif token == "i":
            italic = True
        elif token == "/i":
            italic = False
        elif token == "u":
            underline = True
        elif token == "/u":
            underline = False
        elif token == "s":
            strike = True
        elif token == "/s":
            strike = False
        elif token.startswith("font:"):
            font_family = token.removeprefix("font:")
        elif token == "/font":
            font_family = None
        elif token.startswith("size:"):
            try:
                font_size_px = float(token.removeprefix("size:"))
            except ValueError:
                font_size_px = None
        elif token == "/size":
            font_size_px = None
        cursor = match.end()
    if cursor < len(text):
        runs.append(
            InlineRun(
                text[cursor:],
                bold=bold,
                italic=italic,
                underline=underline,
                strike=strike,
                font_family=font_family,
                font_size_px=font_size_px,
            )
        )
    return [run for run in runs if run.text]


def _apply_inline_run_format(
    run: Any,
    inline: InlineRun,
    *,
    template: Mapping[str, Any] | None,
    profile: Mapping[str, Any],
    block_type: str | None,
) -> None:
    run.bold = bool(template.get("bold", False)) if template else False
    run.italic = bool(template.get("italic", False)) if template else False
    if inline.bold:
        run.bold = True
    if inline.italic:
        run.italic = True
    if inline.underline:
        run.underline = True
    if inline.strike:
        run.font.strike = True
    run.font.name = inline.font_family or str(
        template.get("font_family") if template else profile["docx_translation_font_family"]
    )
    font_size_pt = (
        max(6, min(32, round(inline.font_size_px * 0.65)))
        if inline.font_size_px is not None
        else int(
            template.get("docx_font_size_pt")
            if template
            else _docx_font_size_for_block(profile, block_type=block_type, source=False)
        )
    )
    run.font.size = Pt(font_size_pt)


def _add_translation_paragraphs(
    document: Document,
    text: str,
    *,
    style: str,
    config: TranslationDocxConfig,
    block_type: str | None,
    internal_style_key: str | None = None,
) -> None:
    blocks = _iter_styled_translation_blocks(text, fallback_style_key=internal_style_key)
    if not blocks:
        blocks = [TranslationParagraphBlock(text, None, internal_style_key)]
    for block in blocks:
        block_style = block.style_key or internal_style_key
        paragraph = document.add_paragraph(style=style)
        profile = _effective_style_profile(config)
        template = _style_template(profile, block_style)
        _apply_translation_layout(
            paragraph,
            config,
            block_type=block_type,
            alignment=block.alignment,
            internal_style_key=block_style,
        )
        for inline in _iter_inline_runs(block.text):
            run = paragraph.add_run(inline.text)
            _apply_inline_run_format(
                run,
                inline,
                template=template,
                profile=profile,
                block_type=block_type,
            )


def _heading_level(value: object, default: int = 1) -> int:
    if isinstance(value, int):
        parsed = value
    elif isinstance(value, str):
        try:
            parsed = int(value)
        except ValueError:
            return default
    else:
        return default
    return max(1, min(6, parsed))


def _answer_heading_level(answer: object) -> object:
    if isinstance(answer, Mapping):
        return answer.get("heading_level")
    return None


def _answer_position(answer: object) -> object:
    if isinstance(answer, Mapping):
        return answer.get("position")
    return None


def _answer_display(answer: object) -> bool:
    if isinstance(answer, Mapping) and "display" in answer:
        return bool(answer.get("display"))
    return True


def _toc_position(value: object) -> TocPosition:
    return "back" if value == "back" else "front"


def docx_config_from_pflichtfragen(
    pflichtfragen: Sequence[Mapping[str, Any]] | None,
) -> TranslationDocxConfig:
    """Build a DOCX config from persisted preflight decision-event payloads."""
    values: dict[str, Any] = {}
    for entry in pflichtfragen or []:
        key = entry.get("frage_key")
        answer = entry.get("answer")
        if not isinstance(key, str) or not isinstance(answer, Mapping):
            continue
        values[key] = dict(answer)

    return TranslationDocxConfig(
        header_heading_level=_heading_level(
            _answer_heading_level(values.get("header_heading_level"))
        ),
        chapter_break_heading_level=_heading_level(
            _answer_heading_level(values.get("chapter_break_heading_level"))
        ),
        toc_position=_toc_position(_answer_position(values.get("toc_position"))),
        display_arabic_chapter_headings=_answer_display(
            values.get("display_arabic_chapter_headings")
        ),
    )


def docx_config_from_export_payload(payload: Mapping[str, Any]) -> TranslationDocxConfig:
    export_config = payload.get("export_config")
    if not isinstance(export_config, Mapping):
        return DEFAULT_DOCX_CONFIG
    pflichtfragen = export_config.get("pflichtfragen")
    if not isinstance(pflichtfragen, list):
        base = DEFAULT_DOCX_CONFIG
    else:
        base = docx_config_from_pflichtfragen(pflichtfragen)
    style_profile = export_config.get("style_profile")
    if not isinstance(style_profile, Mapping):
        return base
    return TranslationDocxConfig(
        header_heading_level=base.header_heading_level,
        chapter_break_heading_level=base.chapter_break_heading_level,
        toc_position=base.toc_position,
        display_arabic_chapter_headings=base.display_arabic_chapter_headings,
        style_profile=normalize_style_profile(dict(style_profile)),
    )


def _add_field_run(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _apply_header_style(paragraph: Any, config: TranslationDocxConfig) -> None:
    profile = _effective_style_profile(config)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.font.name = str(profile["docx_translation_font_family"])
        run.font.size = Pt(int(profile["docx_header_font_size_pt"]))


def _set_running_header(
    document: Any, *, project_title: str, config: TranslationDocxConfig
) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_paragraph = (
        first_header.paragraphs[0] if first_header.paragraphs else first_header.add_paragraph()
    )
    first_paragraph.text = project_title
    _apply_header_style(first_paragraph, config)

    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = project_title
    _apply_header_style(paragraph, config)


def _sanitize_headers(document: Any, *, project_title: str, config: TranslationDocxConfig) -> None:
    """Keep headers field-free so Word cannot render stale field errors."""
    for section in document.sections:
        section.different_first_page_header_footer = True
        for header in (section.header, section.first_page_header, section.even_page_header):
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = project_title
            _apply_header_style(paragraph, config)
            for extra in header.paragraphs[1:]:
                extra.text = ""


def _add_toc(document: Any) -> None:
    """Insert a TOC field with `\\o "1-6"` per Formatvorlagen-Baseline
    v1.1 §7.2 (Schluss-Audit Paket 7 Item 2 (a), 2026-05-08). Word
    renders this on first open; python-docx writes it as a field
    instruction.
    """
    paragraph = document.add_paragraph()
    _add_field_run(paragraph, 'TOC \\o "1-6" \\h \\z \\u')


_BLOCK_TYPE_STYLE: dict[str, str] = {
    "main_text": "Normal",
    "MT": "Normal",
    "UE": "Heading 1",
    "HD": "Heading 2",
    "FN": "Footnote Text",
    "QR": "Quote",
    "RN": "Caption",
    "heading": "Heading 1",
    "footnote": "Footnote Text",
    "quran": "Quote",
    "hadith": "Quote",
    "marginalia": "Caption",
}


_DOCX_STYLE_FOR_TRANSLATION_KEY = {
    "body_de": "Normal",
    "body_de_no_indent": "Normal",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "heading_4": "Heading 4",
    "heading_5": "Heading 5",
    "heading_6": "Heading 6",
    "quran_de": "Quote",
    "hadith_de": "Quote",
    "quote_de": "Quote",
    "source_note": "Caption",
    "footnote_text": "Footnote Text",
}


_BLOCK_TYPE_FOR_TRANSLATION_KEY = {
    "heading_1": "UE",
    "heading_2": "HD",
    "heading_3": "heading",
    "heading_4": "heading",
    "heading_5": "heading",
    "heading_6": "heading",
    "quran_de": "quran",
    "hadith_de": "hadith",
    "quote_de": "quote",
    "source_note": "caption",
    "footnote_text": "footnote",
}


def _style_for_block(
    block_type: str,
    config: TranslationDocxConfig,
    internal_style_key: str | None = None,
) -> str:
    if internal_style_key:
        style = _DOCX_STYLE_FOR_TRANSLATION_KEY.get(internal_style_key)
        if style is not None:
            return style
    if block_type == "UE":
        return f"Heading {config.chapter_break_heading_level}"
    if block_type == "HD":
        return f"Heading {min(6, config.chapter_break_heading_level + 1)}"
    return _BLOCK_TYPE_STYLE.get(block_type, "Normal")


def _block_type_for_style_key(block_type: str, internal_style_key: str | None) -> str:
    if not internal_style_key:
        return block_type
    return _BLOCK_TYPE_FOR_TRANSLATION_KEY.get(internal_style_key, block_type)


def _include_source_for_block(block_type: str, config: TranslationDocxConfig) -> bool:
    return config.display_arabic_chapter_headings


def _add_title_page(document: Any, *, project_title: str, config: TranslationDocxConfig) -> None:
    title = document.add_heading(project_title, level=0)
    for run in title.runs:
        run.font.size = Pt(20)
    document.add_paragraph()
    if config.toc_position == "front":
        _add_toc(document)
        document.add_page_break()
    else:
        document.add_page_break()


def _add_back_toc(document: Any, config: TranslationDocxConfig) -> None:
    if config.toc_position != "back":
        return
    document.add_page_break()
    document.add_heading("Contents", level=1)
    _add_toc(document)


async def build_translation_docx(
    *,
    session: AsyncSession,
    project_uuid: _uuid.UUID,
    project_title: str,
    segment_uuids: list[_uuid.UUID] | None = None,
    config: TranslationDocxConfig = DEFAULT_DOCX_CONFIG,
) -> TranslationDocxArtefact:
    """Build the translation-export DOCX in memory.

    Reads only: Page/Block/Segment in-scope of `project_uuid`. The
    artefact captures target-language text per Segment; Arabic-side
    source is preserved on a separate paragraph for the bilingual
    publication layout per Formatvorlagen-Baseline v1.1 §7.2.

    H-4 invariant: this function does not write a Revision row, does
    not write a TRANSLATION-PO row, does not modify Segment text. Pure
    read of `segments.text_content`.
    """
    # Resolve in-scope segments (project-wide if no explicit set).
    stmt = (
        select(Page, Block, Segment)
        .join(Block, Block.page_uuid == Page.page_uuid)
        .join(Segment, Segment.block_uuid == Block.block_uuid)
        .where(Page.project_uuid == project_uuid)
        .where(Segment.active.is_(True))
        .order_by(Page.page_index, Block.block_index, Segment.satz_index)
    )
    if segment_uuids is not None:
        stmt = stmt.where(Segment.satz_uuid.in_(segment_uuids))
    rows = (await session.execute(stmt)).all()
    style_map = await read_translation_style_map(
        session=session,
        segment_uuids=[segment.satz_uuid for _page, _block, segment in rows],
    )

    document = Document()
    _set_document_rtl(document)
    style_profile = _effective_style_profile(config)

    # Page setup per Formatvorlagen-Baseline v1.1 §7.2.
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    _set_running_header(document, project_title=project_title, config=config)
    _add_title_page(document, project_title=project_title, config=config)

    pages_seen: set[_uuid.UUID] = set()
    blocks_seen: set[_uuid.UUID] = set()
    segs_seen: set[_uuid.UUID] = set()
    block_types_seen: set[str] = set()
    last_page_index: int | None = None

    for page, block, segment in rows:
        pages_seen.add(page.page_uuid)
        blocks_seen.add(block.block_uuid)
        segs_seen.add(segment.satz_uuid)
        block_types_seen.add(block.block_type)

        if last_page_index is None or last_page_index != page.page_index:
            heading = document.add_heading(f"Page {page.page_index}", level=1)
            for run in heading.runs:
                run.font.size = Pt(14)
            last_page_index = page.page_index

        text_state = await resolve_segment_text_state(session=session, segment=segment)
        source, target = text_state.source_text, text_state.target_text

        style_key = style_map.get(segment.satz_uuid)
        style = _style_for_block(block.block_type, config, style_key)
        arabic_style = _ensure_arabic_style(document, style)

        if source.strip() and _include_source_for_block(block.block_type, config):
            ar_paragraph = document.add_paragraph(source, style=arabic_style)
            _apply_arabic_layout(ar_paragraph, config, block_type=block.block_type)

        if target.strip() or not source.strip():
            _add_translation_paragraphs(
                document,
                target or "",
                style=style,
                config=config,
                block_type=_block_type_for_style_key(block.block_type, style_key),
                internal_style_key=style_key,
            )

    _add_back_toc(document, config)
    _sanitize_headers(document, project_title=project_title, config=config)

    # Serialize.
    buffer = io.BytesIO()
    document.save(buffer)
    bytes_ = buffer.getvalue()
    sha = hashlib.sha256(bytes_).hexdigest()

    return TranslationDocxArtefact(
        artefact_uuid=new_uuid(),
        bytes_=bytes_,
        sha256=sha,
        size_bytes=len(bytes_),
        n_pages_exported=len(pages_seen),
        n_segments_exported=len(segs_seen),
        block_types_present=sorted(block_types_seen),
        exported_segment_uuids=sorted(segs_seen, key=str),
        exported_page_uuids=sorted(pages_seen, key=str),
        exported_block_uuids=sorted(blocks_seen, key=str),
        protocol={
            "n_pages": len(pages_seen),
            "n_segments": len(segs_seen),
            "block_types_present": sorted(block_types_seen),
            "docx_config": {
                "header_heading_level": config.header_heading_level,
                "chapter_break_heading_level": config.chapter_break_heading_level,
                "toc_position": config.toc_position,
                "display_arabic_chapter_headings": config.display_arabic_chapter_headings,
            },
            "style_profile": style_profile,
        },
    )


async def build_translation_docx_from_snapshot(
    *,
    session: AsyncSession,
    project_uuid: _uuid.UUID,
    project_title: str,
    revision_uuids: list[_uuid.UUID],
    config: TranslationDocxConfig = DEFAULT_DOCX_CONFIG,
) -> TranslationDocxArtefact:
    """Rebuild a translation-export DOCX from a frozen `revision_snapshot[]`.

    Used by the download endpoint to reconstruct the exact text state
    that was exported, even if Segments have been re-translated since.
    Reads `Revision.after_text` (immutable per H-5) for each rev_uuid in
    the snapshot, then assembles the DOCX in Page/Block/satz order using
    the Segments/Blocks/Pages each Revision is anchored to.

    H-4: pure-read; writes nothing.
    """
    if not revision_uuids:
        # Empty snapshot — return a minimal artefact (title + empty body).
        revision_uuids = []
    rows = (
        await session.execute(
            select(Revision, Page, Block, Segment)
            .join(Segment, Segment.satz_uuid == Revision.satz_uuid)
            .join(Block, Block.block_uuid == Segment.block_uuid)
            .join(Page, Page.page_uuid == Block.page_uuid)
            .where(Page.project_uuid == project_uuid)
            .where(Revision.rev_uuid.in_(revision_uuids))
            .order_by(Page.page_index, Block.block_index, Segment.satz_index)
        )
    ).all()
    style_map = await read_translation_style_map(
        session=session,
        segment_uuids=[segment.satz_uuid for _revision, _page, _block, segment in rows],
    )

    document = Document()
    _set_document_rtl(document)
    style_profile = _effective_style_profile(config)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    _set_running_header(document, project_title=project_title, config=config)
    _add_title_page(document, project_title=project_title, config=config)

    pages_seen: set[_uuid.UUID] = set()
    blocks_seen: set[_uuid.UUID] = set()
    segs_seen: set[_uuid.UUID] = set()
    block_types_seen: set[str] = set()
    last_page_index: int | None = None

    for revision, page, block, segment in rows:
        pages_seen.add(page.page_uuid)
        blocks_seen.add(block.block_uuid)
        segs_seen.add(segment.satz_uuid)
        block_types_seen.add(block.block_type)

        if last_page_index is None or last_page_index != page.page_index:
            heading = document.add_heading(f"Page {page.page_index}", level=1)
            for run in heading.runs:
                run.font.size = Pt(14)
            last_page_index = page.page_index

        # Use the Revision's after_text — the immutable snapshot text.
        source = await resolve_segment_source_text(
            session=session,
            segment=segment,
            at_or_before=revision.created_at,
        )
        embedded_source, embedded_target = split_source_target_text(revision.after_text)
        target = embedded_target or revision.after_text
        if embedded_source and not source:
            source = embedded_source
        style_key = style_map.get(segment.satz_uuid)
        style = _style_for_block(block.block_type, config, style_key)
        arabic_style = _ensure_arabic_style(document, style)
        if source.strip() and _include_source_for_block(block.block_type, config):
            ar_paragraph = document.add_paragraph(source, style=arabic_style)
            _apply_arabic_layout(ar_paragraph, config, block_type=block.block_type)
        if target.strip() or not source.strip():
            _add_translation_paragraphs(
                document,
                target or "",
                style=style,
                config=config,
                block_type=_block_type_for_style_key(block.block_type, style_key),
                internal_style_key=style_key,
            )

    _add_back_toc(document, config)
    _sanitize_headers(document, project_title=project_title, config=config)

    buffer = io.BytesIO()
    document.save(buffer)
    bytes_ = buffer.getvalue()
    sha = hashlib.sha256(bytes_).hexdigest()

    return TranslationDocxArtefact(
        artefact_uuid=new_uuid(),
        bytes_=bytes_,
        sha256=sha,
        size_bytes=len(bytes_),
        n_pages_exported=len(pages_seen),
        n_segments_exported=len(segs_seen),
        block_types_present=sorted(block_types_seen),
        exported_segment_uuids=sorted(segs_seen, key=str),
        exported_page_uuids=sorted(pages_seen, key=str),
        exported_block_uuids=sorted(blocks_seen, key=str),
        protocol={
            "n_pages": len(pages_seen),
            "n_segments": len(segs_seen),
            "block_types_present": sorted(block_types_seen),
            "rebuilt_from_snapshot": True,
            "docx_config": {
                "header_heading_level": config.header_heading_level,
                "chapter_break_heading_level": config.chapter_break_heading_level,
                "toc_position": config.toc_position,
                "display_arabic_chapter_headings": config.display_arabic_chapter_headings,
            },
            "style_profile": style_profile,
        },
    )


__all__ = [
    "DEFAULT_DOCX_CONFIG",
    "TranslationDocxArtefact",
    "TranslationDocxConfig",
    "build_translation_docx",
    "build_translation_docx_from_snapshot",
    "docx_config_from_export_payload",
    "docx_config_from_pflichtfragen",
]
