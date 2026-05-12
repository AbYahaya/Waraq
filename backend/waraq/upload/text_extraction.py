"""§2.1 Phase 5 K-2 + K-3 — paragraph extraction for direct-text uploads.

Nine formats covered across two canon rows §2.1:

  K-2 — Document formats:
    - DOCX  — `python-docx`'s `Document(...).paragraphs`
    - ODT   — `odfpy`'s `text:p` elements traversed via the DOM
    - TXT   — plain text split by blank-line paragraph boundaries
    - XML   — text nodes extracted via `xml.etree`, paragraph-split
    - HTML  — text nodes extracted via stdlib `html.parser`, paragraph-split

  K-3 — E-book formats (DjVu is NOT direct-text — it's raster, see
  `page_runner._rasterize_page`'s DjVu branch):
    - EPUB         — `ebooklib`'s spine iteration; each item's XHTML
                     goes through the same block-tag HTML parser used
                     for the bare-HTML path.
    - MOBI / AZW /
      AZW3        — `mobi`'s `extract()` produces an HTML file we
                     parse the same way. AZW/AZW3 are MOBI-family —
                     unless DRM'd, the `mobi` lib reads them.

All nine produce a `list[str]` of non-empty paragraphs in document order.
The upload-service finalize branch materializes one `Segment` per
paragraph, attached to a single `Block` (class MAIN_TEXT) on a single
`Page` per document.

Encoding: TXT/XML/HTML are decoded as UTF-8 by default; on
UnicodeDecodeError we fall back to UTF-8 with `errors='replace'` so
mis-encoded uploads still produce *some* text rather than HTTP 500.
This matches the canon-honest "produce best-effort signal, surface what
went wrong on review" pattern used in the OCR error pipeline.

Empty-paragraph filtering: paragraphs consisting only of whitespace are
dropped. An empty document (no non-whitespace paragraphs) raises
`EmptyDocument` — the service surfaces this as HTTP 422 rather than
silently materializing a Page with zero Segments.

Kindle-DRM honesty: AZW/AZW3 files purchased from the Kindle store are
DRM-protected; `mobi` raises on those. We surface a clear
`TextExtractionError("DRM-protected — cannot extract")` rather than
silently producing garbage or bypassing the DRM.
"""

from __future__ import annotations

import html
import re
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree as ET

from waraq.upload.file_type import UploadFormat

_WHITESPACE_ONLY = re.compile(r"^\s*$")
_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n")  # blank-line paragraph boundary


class TextExtractionError(ValueError):
    """Raised when a direct-text source can't be parsed. The upload
    finalize endpoint converts this to HTTP 422."""


class EmptyDocument(TextExtractionError):
    """Document parsed cleanly but contains no non-whitespace paragraphs.
    Surfaced as HTTP 422 — a 0-Segment Page is a degenerate state we'd
    rather not materialize."""


def extract_paragraphs(*, path: Path, fmt: UploadFormat) -> list[str]:
    """Return the document's paragraphs as a list of non-empty strings.

    Raises:
        TextExtractionError: parser failed or format is not direct-text.
        EmptyDocument: parser succeeded but no non-whitespace text.
    """
    if fmt == UploadFormat.DOCX:
        paragraphs = _extract_docx(path)
    elif fmt == UploadFormat.ODT:
        paragraphs = _extract_odt(path)
    elif fmt == UploadFormat.TXT:
        paragraphs = _extract_txt(path)
    elif fmt == UploadFormat.XML:
        paragraphs = _extract_xml(path)
    elif fmt == UploadFormat.HTML:
        paragraphs = _extract_html(path)
    elif fmt == UploadFormat.EPUB:
        paragraphs = _extract_epub(path)
    elif fmt in (UploadFormat.MOBI, UploadFormat.AZW, UploadFormat.AZW3):
        paragraphs = _extract_mobi(path, fmt)
    else:
        raise TextExtractionError(f"Format {fmt.value!r} is not a direct-text format")

    cleaned = [p.strip() for p in paragraphs if not _WHITESPACE_ONLY.match(p or "")]
    if not cleaned:
        raise EmptyDocument(
            f"Document at {path} contains no non-whitespace paragraphs"
        )
    return cleaned


# ---------------------------------------------------------------------
# DOCX — python-docx
# ---------------------------------------------------------------------


def _extract_docx(path: Path) -> list[str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise TextExtractionError(
            "python-docx not installed. Add `python-docx` to the venv."
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise TextExtractionError(f"python-docx could not open {path}: {exc!r}") from exc
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------
# ODT — odfpy
# ---------------------------------------------------------------------


def _extract_odt(path: Path) -> list[str]:
    try:
        from odf.opendocument import load
        from odf.text import P
    except ImportError as exc:
        raise TextExtractionError("odfpy not installed. Add `odfpy` to the venv.") from exc

    try:
        doc = load(str(path))
    except Exception as exc:
        raise TextExtractionError(f"odfpy could not open {path}: {exc!r}") from exc

    paragraphs: list[str] = []
    for p in doc.getElementsByType(P):
        # `teletype.extractText` is the canonical odfpy text helper, but
        # it's a small enough operation to inline so we don't depend on
        # an import path that's renamed across odfpy versions.
        text_parts: list[str] = []
        _collect_text_from_node(p, text_parts)
        paragraphs.append("".join(text_parts))
    return paragraphs


def _collect_text_from_node(node: object, out: list[str]) -> None:
    """Recursively collect text from an ODF DOM node."""
    data = getattr(node, "data", None)
    if isinstance(data, str):
        out.append(data)
    children = getattr(node, "childNodes", ()) or ()
    for child in children:
        _collect_text_from_node(child, out)


# ---------------------------------------------------------------------
# TXT — split by blank-line paragraph boundary
# ---------------------------------------------------------------------


def _extract_txt(path: Path) -> list[str]:
    raw = _read_text_with_fallback(path)
    # Normalize CRLF / CR → LF, then split on blank lines.
    normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
    return _PARAGRAPH_SPLIT.split(normalized)


# ---------------------------------------------------------------------
# XML — extract text nodes, then paragraph-split
# ---------------------------------------------------------------------


def _extract_xml(path: Path) -> list[str]:
    raw_text = _read_text_with_fallback(path)
    try:
        root = ET.fromstring(raw_text)
    except ET.ParseError as exc:
        raise TextExtractionError(f"XML parse failed for {path}: {exc!r}") from exc

    # Walk the tree, collect all text + tail nodes in document order.
    pieces: list[str] = []
    for elem in root.iter():
        if elem.text:
            pieces.append(elem.text)
        if elem.tail:
            pieces.append(elem.tail)
    flat = "\n\n".join(s for s in (p.strip() for p in pieces) if s)
    return _PARAGRAPH_SPLIT.split(flat)


# ---------------------------------------------------------------------
# HTML — stdlib html.parser, strip tags, paragraph-split
# ---------------------------------------------------------------------


class _HtmlTextCollector(HTMLParser):
    """Collect text from `<p>`, `<div>`, `<li>`, `<h1>...<h6>` etc.
    Each block-level close becomes a paragraph boundary. Inline elements
    (`<span>`, `<b>`, etc.) flow inside the current paragraph."""

    _BLOCK_TAGS = frozenset(
        {
            "p",
            "div",
            "li",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "blockquote",
            "section",
            "article",
            "header",
            "footer",
            "main",
            "nav",
            "aside",
            "br",
            "tr",
        }
    )
    _SKIP_TAGS = frozenset({"script", "style", "head"})

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._paragraphs: list[list[str]] = [[]]
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        elif tag in self._BLOCK_TAGS and self._paragraphs[-1]:
            self._paragraphs.append([])

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth = max(self._skip_depth - 1, 0)
        elif tag in self._BLOCK_TAGS and self._paragraphs[-1]:
            self._paragraphs.append([])

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if data.strip():
            self._paragraphs[-1].append(data)

    def paragraphs(self) -> list[str]:
        return ["".join(parts) for parts in self._paragraphs if parts]


def _extract_html(path: Path) -> list[str]:
    raw = _read_text_with_fallback(path)
    return _parse_html_string(raw)


def _parse_html_string(raw: str) -> list[str]:
    """Block-tag-aware paragraph extraction from an HTML string. Shared
    between the bare-HTML upload path (K-2) and the EPUB-spine /
    MOBI-extract paths (K-3) — every K-3 e-book format ultimately
    decomposes into HTML chunks that get parsed the same way."""
    parser = _HtmlTextCollector()
    try:
        parser.feed(raw)
        parser.close()
    except Exception as exc:
        raise TextExtractionError(f"HTML parse failed: {exc!r}") from exc
    # Unescape any leftover entities that `convert_charrefs` missed
    # (very old HTML may have malformed entities that slip through).
    return [html.unescape(p) for p in parser.paragraphs()]


# ---------------------------------------------------------------------
# EPUB — ebooklib spine iteration → per-item HTML → paragraphs
# ---------------------------------------------------------------------


def _extract_epub(path: Path) -> list[str]:
    """Walk the EPUB spine in reading order, extract each item's
    XHTML body, and concatenate paragraphs. The spine is the canonical
    "linear reading order" of an EPUB per the spec; iterating in spine
    order matches what the user sees in an EPUB reader."""
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError as exc:
        raise TextExtractionError("ebooklib not installed. Add `ebooklib` to the venv.") from exc

    try:
        book = epub.read_epub(str(path))
    except Exception as exc:
        raise TextExtractionError(f"ebooklib could not open EPUB {path}: {exc!r}") from exc

    out: list[str] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        # Skip navigation scaffolding (TOC / nav docs) — they list
        # chapter titles, not body content. Three exclusion paths to
        # cover ebooklib version variance: (a) EPUB-3 `properties=['nav']`
        # marker, (b) instance check against `epub.EpubNav`, (c) NCX
        # media-type for legacy EPUB-2 TOCs.
        properties = getattr(item, "properties", None) or []
        if "nav" in properties:
            continue
        if isinstance(item, epub.EpubNav):
            continue
        media_type = getattr(item, "media_type", "") or ""
        if "ncx" in media_type.lower():
            continue
        try:
            raw = item.get_content()
        except Exception:
            continue
        if isinstance(raw, bytes):
            try:
                raw_text = raw.decode("utf-8")
            except UnicodeDecodeError:
                raw_text = raw.decode("utf-8", errors="replace")
        else:
            raw_text = str(raw)
        out.extend(_parse_html_string(raw_text))
    return out


# ---------------------------------------------------------------------
# MOBI / AZW / AZW3 — `mobi.extract()` → tmpdir of HTML → paragraphs
# ---------------------------------------------------------------------


def _extract_mobi(path: Path, fmt: UploadFormat) -> list[str]:
    """`mobi.extract(path)` writes a temp directory containing one or
    more `.html` files (depending on the format). We concatenate the
    HTML files in lexical order (Calibre / mobitool's natural sort)
    and parse via the shared HTML extractor.

    DRM-protected files raise an opaque error inside `mobi.extract` —
    we wrap it as `TextExtractionError("DRM-protected — cannot extract")`
    rather than bypassing the DRM (which would be a §7.4 violation).
    """
    try:
        import mobi
    except ImportError as exc:
        raise TextExtractionError("mobi not installed. Add `mobi` to the venv.") from exc

    try:
        tempdir, filepath = mobi.extract(str(path))
    except Exception as exc:
        msg = repr(exc).lower()
        if "drm" in msg or "encrypted" in msg:
            raise TextExtractionError(
                f"{fmt.value.upper()} file appears DRM-protected — extraction refused. "
                "Remove DRM (legally) before uploading."
            ) from exc
        raise TextExtractionError(
            f"mobi could not extract {fmt.value.upper()} from {path}: {exc!r}"
        ) from exc

    out: list[str] = []
    try:
        # `mobi.extract` returns a tuple `(tempdir, primary_file)`.
        # The primary file is the entry point; other HTML files may
        # exist for split MOBI/AZW3 documents. Collect every `.html`
        # file in lexical order under the tempdir.
        from pathlib import Path as _Path

        root = _Path(tempdir)
        html_files = sorted(root.rglob("*.html")) + sorted(root.rglob("*.htm"))
        if not html_files and _Path(filepath).exists():
            html_files = [_Path(filepath)]

        for hf in html_files:
            try:
                raw = hf.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                raw = hf.read_text(encoding="utf-8", errors="replace")
            out.extend(_parse_html_string(raw))
    finally:
        # mobi.extract leaves a tempdir behind; clean up.
        try:
            import shutil

            shutil.rmtree(tempdir, ignore_errors=True)
        except Exception:
            pass

    return out


# ---------------------------------------------------------------------
# Shared text-read helper
# ---------------------------------------------------------------------


def _read_text_with_fallback(path: Path) -> str:
    """Read a text file as UTF-8; on decode error fall back to
    `errors='replace'` so the upload still produces some text. The
    OCR-stage error pipeline pattern: produce best-effort signal,
    surface what went wrong on review."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


__all__ = [
    "EmptyDocument",
    "TextExtractionError",
    "extract_paragraphs",
]
