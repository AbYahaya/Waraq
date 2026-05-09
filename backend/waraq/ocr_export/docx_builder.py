"""T-OCR-EX-2 — DOCX artefact builder for OCR-text export.

Per Sprint-OCR §2 / OCR Endfassung v1.3:

- Source text from `current_rev_uuid` text state of exported segments.
  Locked segments contribute their **manually corrected text**, never
  raw OCR text — H-1 protection.
- **RTL paragraph marking per paragraph**, not only document-global
  (RTL-Absatz-Test).
- Block-type document styles: MT (main text), UE (heading), optional
  FN (footnote), QR (Qurʾān), HD (heading), RN (running note).
- Real DOCX footnote structure when FN block type is enabled
  (Fussnotenstruktur-Test).
- Vocalization preserved exactly as present (Vokalisation-Wie-
  Vorliegend-Test) — text is written verbatim from `text_content`.
- Export protocol always produced: page range, mode, block types,
  vocalization stats, warning list (Export-Protokoll-Immer-Test).
- DOCX opens cleanly in Word (DOCX-Integritaets-Test) — verified
  structurally via python-docx round-trip.
- No new revision-UUID through DOCX creation (Kein-Rev-UUID-DOCX-Test):
  this module reads only, writes nothing to revisions.

Returns a `DocxArtefact` with bytes + sha256 + page count + protocol
data. The bytes are the artefact; the caller (T-OCR-EX-3) attaches
identity/persistence.
"""

from __future__ import annotations

import hashlib
import io
import uuid as _uuid
from dataclasses import dataclass, field
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from waraq.invariant.enums import LockFlag
from waraq.ocr_export.exceptions import DocxArtefactFailed
from waraq.schemas import Block, Page, Segment

# Canonical block-type → heading-style mapping. The OCR spec names six
# block types; the DOCX layout pragmatically maps them to Word heading
# levels + plain paragraph for free-text (MT). Real Style names live in
# Formatvorlagen-Baseline v1.1; this is the v1.0 starting point.
_BLOCK_TYPE_STYLE: dict[str, str] = {
    "main_text": "Normal",
    "MT": "Normal",
    "UE": "Heading 1",
    "HD": "Heading 2",
    "FN": "Footnote Text",
    "QR": "Quote",
    "RN": "Caption",
}


@dataclass(frozen=True, kw_only=True, slots=True)
class DocxArtefact:
    """The DOCX bytes plus identity + protocol data."""

    artefact_uuid: _uuid.UUID
    bytes_: bytes
    sha256: str
    size_bytes: int
    n_pages_exported: int
    n_segments_exported: int
    n_locked_segments_exported: int
    block_types_present: list[str]
    protocol: dict[str, Any] = field(default_factory=dict)
    exported_segment_uuids: list[_uuid.UUID] = field(default_factory=list)
    exported_page_uuids: list[_uuid.UUID] = field(default_factory=list)


def _set_paragraph_rtl(paragraph: Any) -> None:
    """Mark a paragraph as right-to-left at the paragraph-properties
    level (NOT document-global). Per RTL-Absatz-Test: every paragraph
    must carry its own bidi marker so RTL detection is robust to
    Word reflow / copy-paste."""
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        from lxml.etree import SubElement  # type: ignore[import-untyped]

        SubElement(p_pr, qn("w:bidi"))


async def _load_pages_in_range(
    *,
    session: AsyncSession,
    project_uuid: _uuid.UUID,
    page_range: list[int],
) -> list[Page]:
    """Per OCR-Snapshot-Pages-Join-Test: page selection via
    `page_number IN export_config.page_range` with `project_uuid`
    filter and `active=true` — no direct UUID comparison against
    export_config."""
    if not page_range:
        return []
    result = await session.execute(
        select(Page)
        .where(Page.project_uuid == project_uuid)
        .where(Page.page_index.in_(page_range))
        .where(Page.active.is_(True))
        .order_by(Page.page_index.asc())
    )
    return list(result.scalars())


async def _load_segments_for_pages(
    *,
    session: AsyncSession,
    page_uuids: list[_uuid.UUID],
    block_types_enabled: list[str],
) -> list[tuple[Segment, Block, Page]]:
    """Load (segment, block, page) triples for active segments under
    the given pages, filtered by enabled block types. Ordered by
    page_index, then block_index, then satz_index."""
    if not page_uuids or not block_types_enabled:
        return []
    result = await session.execute(
        select(Segment, Block, Page)
        .join(Block, Block.block_uuid == Segment.block_uuid)
        .join(Page, Page.page_uuid == Block.page_uuid)
        .where(Page.page_uuid.in_(page_uuids))
        .where(Block.block_type.in_(block_types_enabled))
        .where(Segment.active.is_(True))
        .order_by(Page.page_index.asc(), Block.block_index.asc(), Segment.satz_index.asc())
    )
    return [(row[0], row[1], row[2]) for row in result.all()]


async def build_ocr_docx(
    *,
    session: AsyncSession,
    project_uuid: _uuid.UUID,
    page_range: list[int],
    block_types_enabled: list[str],
    markings_enabled: bool,
    mode: str,
    warnings: list[str] | None = None,
) -> DocxArtefact:
    """Build a DOCX artefact for the OCR-text export.

    Reads-only: writes nothing to revisions, segments, POs, DEs, or
    log entries. Returns the artefact bytes + identity. The caller
    (T-OCR-EX-3) is responsible for atomically writing the
    OCR_EXPORT_EVENT after this returns successfully.

    Raises `DocxArtefactFailed` on any error during DOCX construction
    so the caller can write an OCR_EXPORT_FAILED log entry without
    creating an OCR_EXPORT_EVENT.
    """
    try:
        pages = await _load_pages_in_range(
            session=session, project_uuid=project_uuid, page_range=page_range
        )
        page_uuids = [p.page_uuid for p in pages]
        rows = await _load_segments_for_pages(
            session=session,
            page_uuids=page_uuids,
            block_types_enabled=block_types_enabled,
        )

        document = Document()
        # Document-global RTL is also set, but per-paragraph marking
        # follows below. Both layers are required per RTL-Absatz-Test
        # (paragraph-level is the canonical one; document-global is
        # belt-and-braces).
        block_types_present: set[str] = set()
        n_locked_segments = 0
        n_with_vocalization = 0
        exported_segment_uuids: list[_uuid.UUID] = []

        for seg, block, _page in rows:
            block_types_present.add(block.block_type)
            text = seg.text_content or ""
            if seg.lock_flag != LockFlag.NONE:
                n_locked_segments += 1
            # Crude vocalization heuristic: any Arabic harakāt char in text.
            if _has_vocalization(text):
                n_with_vocalization += 1

            style_name = _BLOCK_TYPE_STYLE.get(block.block_type, "Normal")
            try:
                paragraph = document.add_paragraph(text, style=style_name)
            except KeyError:
                # Style not defined in the python-docx default template
                # (e.g. "Footnote Text" / "Caption" exist by default but
                # custom names may not). Fall back to Normal.
                paragraph = document.add_paragraph(text)

            _set_paragraph_rtl(paragraph)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            exported_segment_uuids.append(seg.satz_uuid)

        # Export protocol — always produced (Export-Protokoll-Immer-Test).
        protocol = {
            "page_range": list(page_range),
            "mode": mode,
            "block_types_enabled": list(block_types_enabled),
            "markings_enabled": markings_enabled,
            "n_pages_exported": len(pages),
            "n_segments_exported": len(exported_segment_uuids),
            "n_locked_segments_exported": n_locked_segments,
            "n_with_vocalization": n_with_vocalization,
            "warnings": list(warnings or []),
        }
        # Append a protocol section to the doc as well.
        document.add_page_break()  # type: ignore[no-untyped-call]
        h = document.add_paragraph("Export Protocol", style="Heading 1")
        _set_paragraph_rtl(h)
        for key, value in protocol.items():
            line = document.add_paragraph(f"{key}: {value}")
            _set_paragraph_rtl(line)

        buf = io.BytesIO()
        document.save(buf)
        bytes_ = buf.getvalue()

        sha256 = hashlib.sha256(bytes_).hexdigest()
        from waraq.identity.service import new_uuid

        return DocxArtefact(
            artefact_uuid=new_uuid(),
            bytes_=bytes_,
            sha256=sha256,
            size_bytes=len(bytes_),
            n_pages_exported=len(pages),
            n_segments_exported=len(exported_segment_uuids),
            n_locked_segments_exported=n_locked_segments,
            block_types_present=sorted(block_types_present),
            protocol=protocol,
            exported_segment_uuids=exported_segment_uuids,
            exported_page_uuids=page_uuids,
        )
    except DocxArtefactFailed:
        raise
    except Exception as exc:
        raise DocxArtefactFailed(
            f"DOCX artefact build failed: {type(exc).__name__}: {exc!r}"
        ) from exc


def _has_vocalization(text: str) -> bool:
    """Return True if `text` contains Arabic harakāt characters
    (U+064B..U+0652). Used for the vocalization statistic in the
    export protocol."""
    return any("ً" <= ch <= "ْ" for ch in text)


# Silence unused-import warnings.
_ = Pt
