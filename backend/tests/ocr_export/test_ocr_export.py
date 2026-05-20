"""Sprint-OCR — T-OCR-EX-1 + T-OCR-EX-2 + T-OCR-EX-3 tests.

Mandatory tests from OCR Endfassung v1.3 §4 (selected mandatory subset
appropriate to the v1.0 implementation scope):

- OCR-Gate-Blockiert-F06-Test
- OCR-Gate-Blockiert-F07-Test
- OCR-Gate-Blockiert-F08-Test
- OCR-Gate-Vorabpruefung-Kein-Log-Test
- OCR-Gate-Blockiert-Start-Kein-Log-Test
- OCR-Gate-Pflichtfragen-Aktiv-Test
- OCR-Gate-Kein-Profil-Bypass-Test
- OCR-Gate-Decision-Event-Source-Test
- OCR-Gate-Export-Attempt-ID-Test
- RTL-Absatz-Test
- DOCX-Integritaets-Test
- Blocktypen-Filter-Test
- Vokalisation-Wie-Vorliegend-Test
- Gesperrtes-Segment-Manueller-Text-Test
- Export-Protokoll-Immer-Test
- Kein-Rev-UUID-DOCX-Test
- OCR-EXPORT_EVENT-Nur-Bei-Erfolg-Test
- OCR-EXPORT_EVENT-Kein-Eintrag-Bei-Fehler-Test
- OCR-EXPORT_EVENT-Atomaritaet-Test
- OCR-Snapshot-Vollstaendigkeit-Test
- OCR-Decision-Snapshot-Allowlist-Test
- OCR-Decision-Snapshot-Attempt-Bindung-Test
- OCR-Gate-Mode-Test
- OCR-EXPORT_EVENT-Via-PROVENANCE-Kern-Test
- Log-Eintrag-Bei-Gestarteten-Job-Test
"""

from __future__ import annotations

import io

import pytest
from docx import Document  # type: ignore[import-untyped]
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from waraq.decisions import create_decision_event
from waraq.identity import new_uuid
from waraq.invariant.enums import LockFlag, OperationMode
from waraq.ocr.error_classes import OcrErrorClass
from waraq.ocr.review import record_ocr_error_instance
from waraq.ocr_export import (
    DocxArtefactFailed,
    GateMode,
    OcrExportBlocked,
    OcrExportConfig,
    OcrExportGateState,
    OcrExportPflichtfragenMissing,
    Pflichtfragen,
    build_ocr_docx,
    check_ocr_export_gate,
    confirm_pflichtfragen,
    run_ocr_export,
)
from waraq.revision import create_revision
from waraq.schemas import (
    Block,
    DecisionEvent,
    LogEntry,
    Page,
    Project,
    ProvenanceObject,
    Revision,
    Segment,
)
from waraq.schemas.enums import ChangeSource, DecisionSource, OcrStatus, POType, ScopeType


async def _seed_project(
    session: AsyncSession,
    *,
    n_segments: int = 2,
    block_type: str = "MT",
    locked_segment_text: str | None = None,
    initial_text: str = "نص عربي",
) -> tuple[Project, Page, Block, list[Segment]]:
    """Seed a project + 1 page + 1 block + N segments. The 0th segment can
    optionally be locked with custom text (for the 'manual_local must
    write manually corrected text' check)."""
    from tests.conftest import seed_account_uuid

    account_uuid = new_uuid()
    await seed_account_uuid(session, account_uuid)

    project = Project(project_uuid=new_uuid(), account_uuid=account_uuid, name="ocr-export-test")
    session.add(project)
    await session.flush()

    page = Page(
        page_uuid=new_uuid(),
        project_uuid=project.project_uuid,
        page_index=1,
        ocr_status=OcrStatus.GO,
    )
    session.add(page)
    await session.flush()

    block = Block(
        block_uuid=new_uuid(),
        page_uuid=page.page_uuid,
        block_type=block_type,
        block_index=1,
    )
    session.add(block)
    await session.flush()

    segments: list[Segment] = []
    for i in range(n_segments):
        text = (
            locked_segment_text
            if (i == 0 and locked_segment_text is not None)
            else f"{initial_text} {i}"
        )
        lock = (
            LockFlag.MANUAL_LOCAL if (i == 0 and locked_segment_text is not None) else LockFlag.NONE
        )
        seg = Segment(
            satz_uuid=new_uuid(),
            block_uuid=block.block_uuid,
            satz_index=i + 1,
            lock_flag=lock,
            text_content=text,
        )
        session.add(seg)
        segments.append(seg)
    await session.flush()
    return project, page, block, segments


def _basic_pflichtfragen(
    *,
    block_types: list[str] | None = None,
    mode: GateMode = GateMode.ARBEITSSTAND,
) -> Pflichtfragen:
    return Pflichtfragen(
        page_range=[1],
        block_types_enabled=block_types if block_types is not None else ["MT"],
        markings_enabled=True,
        mode=mode,
    )


# --- T-OCR-EX-1: gate hard-block tests --------------------------


@pytest.mark.asyncio
class TestGateHardBlocks:
    """OCR-Gate-Blockiert-F06-Test / -F07-Test / -F08-Test."""

    @pytest.mark.parametrize(
        "code",
        [OcrErrorClass.F_06_QR, OcrErrorClass.F_07, OcrErrorClass.F_08],
    )
    async def test_unresolved_blocking_f_code_blocks_gate(
        self, db_session: AsyncSession, code: OcrErrorClass
    ) -> None:
        project, page, _, _ = await _seed_project(db_session)
        await record_ocr_error_instance(
            session=db_session, page_uuid=page.page_uuid, error_code=code
        )
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=str(new_uuid()),
        )
        result = await check_ocr_export_gate(session=db_session, config=config)
        assert result.state == OcrExportGateState.BLOCKIERT
        assert any(code.value in r for r in result.blocking_reasons)


# --- T-OCR-EX-1: pre-check writes no log -----------------------


@pytest.mark.asyncio
class TestGatePreCheckWritesNoLog:
    """OCR-Gate-Vorabpruefung-Kein-Log-Test +
    OCR-Gate-Blockiert-Start-Kein-Log-Test."""

    async def test_check_ocr_export_gate_writes_no_log_entry(
        self, db_session: AsyncSession
    ) -> None:
        project, _, _, _ = await _seed_project(db_session)
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=str(new_uuid()),
        )
        before = (await db_session.execute(select(func.count()).select_from(LogEntry))).scalar_one()
        await check_ocr_export_gate(session=db_session, config=config)
        after = (await db_session.execute(select(func.count()).select_from(LogEntry))).scalar_one()
        assert after == before

    async def test_run_ocr_export_on_blocked_gate_writes_no_log_entry_and_no_job(
        self, db_session: AsyncSession
    ) -> None:
        project, page, _, _ = await _seed_project(db_session)
        # Inject F-08 to block the gate.
        await record_ocr_error_instance(
            session=db_session, page_uuid=page.page_uuid, error_code=OcrErrorClass.F_08
        )
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=str(new_uuid()),
        )
        before_logs = (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one()
        from waraq.schemas import Job

        before_jobs = (await db_session.execute(select(func.count()).select_from(Job))).scalar_one()

        with pytest.raises(OcrExportBlocked):
            await run_ocr_export(session=db_session, config=config)

        after_logs = (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one()
        after_jobs = (await db_session.execute(select(func.count()).select_from(Job))).scalar_one()
        assert after_logs == before_logs
        assert after_jobs == before_jobs


# --- T-OCR-EX-1: Pflichtfragen ---------------------------------


@pytest.mark.asyncio
class TestPflichtfragen:
    """OCR-Gate-Pflichtfragen-Aktiv-Test +
    OCR-Gate-Decision-Event-Source-Test +
    OCR-Gate-Export-Attempt-ID-Test."""

    async def test_empty_page_range_blocks_gate(self, db_session: AsyncSession) -> None:
        project, _, _, _ = await _seed_project(db_session)
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=Pflichtfragen(
                page_range=[],
                block_types_enabled=["MT"],
                markings_enabled=True,
                mode=GateMode.ARBEITSSTAND,
            ),
            export_attempt_id=str(new_uuid()),
        )
        result = await check_ocr_export_gate(session=db_session, config=config)
        assert result.state == OcrExportGateState.BLOCKIERT

    async def test_empty_block_types_blocks_gate(self, db_session: AsyncSession) -> None:
        project, _, _, _ = await _seed_project(db_session)
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=Pflichtfragen(
                page_range=[1],
                block_types_enabled=[],
                markings_enabled=True,
                mode=GateMode.ARBEITSSTAND,
            ),
            export_attempt_id=str(new_uuid()),
        )
        result = await check_ocr_export_gate(session=db_session, config=config)
        assert result.state == OcrExportGateState.BLOCKIERT

    async def test_confirm_pflichtfragen_writes_de_with_export_confirmation_source(
        self, db_session: AsyncSession
    ) -> None:
        project, _, _, _ = await _seed_project(db_session)
        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=attempt_id,
        )

        de = await confirm_pflichtfragen(session=db_session, config=config)

        assert str(de.decision_source) == DecisionSource.EXPORT_CONFIRMATION.value
        assert de.related_export_attempt_id == attempt_id
        assert str(de.scope_type) == ScopeType.PROJECT.value
        assert de.scope_uuid == project.project_uuid
        assert de.decision_type == "ocr_export_pflichtfragen_bestaetigt"

    async def test_confirm_pflichtfragen_refuses_empty_attempt_id(
        self, db_session: AsyncSession
    ) -> None:
        project, _, _, _ = await _seed_project(db_session)
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id="",
        )
        with pytest.raises(ValueError, match="non-empty export_attempt_id"):
            await confirm_pflichtfragen(session=db_session, config=config)

    async def test_run_ocr_export_refuses_without_pflichtfragen_de(
        self, db_session: AsyncSession
    ) -> None:
        project, _, _, _ = await _seed_project(db_session)
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=str(new_uuid()),
        )
        # No confirm_pflichtfragen call → no DE → run should refuse.
        with pytest.raises(OcrExportPflichtfragenMissing):
            await run_ocr_export(session=db_session, config=config)


# --- T-OCR-EX-2: DOCX builder ----------------------------------


@pytest.mark.asyncio
class TestDocxBuilder:
    async def test_rtl_paragraph_marking_per_paragraph(self, db_session: AsyncSession) -> None:
        """RTL-Absatz-Test: each paragraph carries explicit bidi marker."""
        project, _, _, _ = await _seed_project(db_session, n_segments=2)
        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        # Re-open the produced bytes and inspect every paragraph.
        doc = Document(io.BytesIO(artefact.bytes_))
        from docx.oxml.ns import qn  # type: ignore[import-untyped]

        # At least one body paragraph; every one with non-empty text has
        # the bidi element. (The Heading paragraph for Export Protocol
        # also carries it since we set RTL on it.)
        any_text_paragraphs = [p for p in doc.paragraphs if p.text]
        assert any_text_paragraphs
        for p in any_text_paragraphs:
            p_pr = p._p.find(qn("w:pPr"))
            assert p_pr is not None, f"paragraph missing pPr: {p.text!r}"
            assert p_pr.find(qn("w:bidi")) is not None, f"paragraph missing bidi marker: {p.text!r}"

    async def test_docx_integrity_round_trip_via_python_docx(
        self, db_session: AsyncSession
    ) -> None:
        """DOCX-Integritaets-Test (proxy via python-docx round-trip):
        re-opening the produced bytes succeeds without errors."""
        project, _, _, _ = await _seed_project(db_session, n_segments=2)
        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        doc = Document(io.BytesIO(artefact.bytes_))
        # Sanity: the body has paragraphs.
        assert len(doc.paragraphs) > 0

    async def test_block_type_filter_excludes_disabled_types(
        self, db_session: AsyncSession
    ) -> None:
        """Blocktypen-Filter-Test: only enabled block_types appear; the
        protocol records which types were enabled."""
        # Seed two segments under MT block...
        project, page, _, _mt_segs = await _seed_project(
            db_session, n_segments=2, block_type="MT", initial_text="MT_TEXT"
        )
        # ...then add a QR block + segment under the same page.
        qr_block = Block(
            block_uuid=new_uuid(),
            page_uuid=page.page_uuid,
            block_type="QR",
            block_index=2,
        )
        db_session.add(qr_block)
        await db_session.flush()
        qr_seg = Segment(
            satz_uuid=new_uuid(),
            block_uuid=qr_block.block_uuid,
            satz_index=1,
            lock_flag=LockFlag.NONE,
            text_content="QR_VERSE_TEXT",
        )
        db_session.add(qr_seg)
        await db_session.flush()

        # Build DOCX with QR DISABLED.
        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],  # QR not enabled
            markings_enabled=False,
            mode="arbeitsstand",
        )
        doc = Document(io.BytesIO(artefact.bytes_))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # MT segments rendered.
        assert "MT_TEXT" in all_text
        # QR text NOT rendered.
        assert "QR_VERSE_TEXT" not in all_text
        # Protocol records the enabled set.
        assert artefact.protocol["block_types_enabled"] == ["MT"]
        # `block_types_present` reflects what actually rendered.
        assert "MT" in artefact.block_types_present
        assert "QR" not in artefact.block_types_present

    async def test_vocalization_preserved_as_present(self, db_session: AsyncSession) -> None:
        """Vokalisation-Wie-Vorliegend-Test: harakāt characters are
        carried verbatim (no addition, no suppression)."""
        # Seed segments: one with harakāt (fatha = U+064E), one without.
        vocalized = "كَتَبَ"  # ka-ta-ba with fatha
        unvocalized = "كتب"
        project, _page, block, _ = await _seed_project(db_session, n_segments=0, block_type="MT")
        for i, txt in enumerate([vocalized, unvocalized]):
            db_session.add(
                Segment(
                    satz_uuid=new_uuid(),
                    block_uuid=block.block_uuid,
                    satz_index=i + 1,
                    lock_flag=LockFlag.NONE,
                    text_content=txt,
                )
            )
        await db_session.flush()

        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        doc = Document(io.BytesIO(artefact.bytes_))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert vocalized in all_text  # harakāt preserved
        assert unvocalized in all_text  # un-harakāt also preserved verbatim
        # Protocol's vocalization stat counts the harakāt segment exactly.
        assert artefact.protocol["n_with_vocalization"] == 1

    async def test_locked_segment_writes_manual_text_not_raw_ocr(
        self, db_session: AsyncSession
    ) -> None:
        """Gesperrtes-Segment-Manueller-Text-Test: locked segments
        contribute the (manually corrected) text_content. Since
        text_content for a locked Segment IS the manual correction
        (the H-1 protection ensures the OCR pipeline never overwrote
        it), this is upheld by reading text_content as-is."""
        manual_text = "MANUAL_CORRECTION_HERE"
        project, _, _, segments = await _seed_project(
            db_session, n_segments=2, locked_segment_text=manual_text
        )
        assert segments[0].lock_flag == LockFlag.MANUAL_LOCAL
        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        doc = Document(io.BytesIO(artefact.bytes_))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert manual_text in all_text
        assert artefact.n_locked_segments_exported == 1

    async def test_protocol_always_produced_with_required_fields(
        self, db_session: AsyncSession
    ) -> None:
        """Export-Protokoll-Immer-Test: protocol always present with
        page_range, mode, block_types, vocalization stats, warning list."""
        project, _, _, _ = await _seed_project(db_session, n_segments=1)
        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
            warnings=["sample warning"],
        )
        p = artefact.protocol
        for key in (
            "page_range",
            "mode",
            "block_types_enabled",
            "markings_enabled",
            "n_pages_exported",
            "n_segments_exported",
            "n_with_vocalization",
            "warnings",
        ):
            assert key in p
        assert p["warnings"] == ["sample warning"]

    async def test_docx_creation_writes_no_revision(self, db_session: AsyncSession) -> None:
        """Kein-Rev-UUID-DOCX-Test."""
        project, _, _, _ = await _seed_project(db_session, n_segments=2)
        before = (await db_session.execute(select(func.count()).select_from(Revision))).scalar_one()
        await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        after = (await db_session.execute(select(func.count()).select_from(Revision))).scalar_one()
        assert after == before

    async def test_translation_revision_does_not_replace_ocr_text_in_export(
        self, db_session: AsyncSession
    ) -> None:
        """OCR export must continue to use the latest non-translation
        revision as its source text, even after a translation pass has
        updated `segments.text_content`."""
        project, _, _, segments = await _seed_project(
            db_session,
            n_segments=1,
            initial_text="RAW_SEED",
        )
        segment = segments[0]

        await create_revision(
            session=db_session,
            segment=segment,
            after_text="نص عربي مصحح",
            change_source=ChangeSource.OCR,
            operation_mode=OperationMode.AUTOMATIC,
        )
        await create_revision(
            session=db_session,
            segment=segment,
            after_text="German translated output",
            change_source=ChangeSource.RE_TRANSLATE,
            operation_mode=OperationMode.AUTOMATIC,
        )

        artefact = await build_ocr_docx(
            session=db_session,
            project_uuid=project.project_uuid,
            page_range=[1],
            block_types_enabled=["MT"],
            markings_enabled=True,
            mode="arbeitsstand",
        )
        doc = Document(io.BytesIO(artefact.bytes_))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "نص عربي مصحح" in all_text
        assert "German translated output" not in all_text


# --- T-OCR-EX-3: OCR_EXPORT_EVENT atomicity --------------------


@pytest.mark.asyncio
class TestOcrExportEventAtomicity:
    async def test_full_run_produces_single_ocr_export_event_po(
        self, db_session: AsyncSession
    ) -> None:
        """OCR-EXPORT_EVENT-Nur-Bei-Erfolg-Test +
        OCR-EXPORT_EVENT-Atomaritaet-Test +
        OCR-EXPORT_EVENT-Via-PROVENANCE-Kern-Test +
        Log-Eintrag-Bei-Gestarteten-Job-Test."""
        project, _, _, _ = await _seed_project(db_session, n_segments=2)
        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=attempt_id,
        )
        await confirm_pflichtfragen(session=db_session, config=config)

        before_pos = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.OCR_EXPORT_EVENT.value)
            )
        ).scalar_one()
        before_logs = (
            await db_session.execute(
                select(func.count())
                .select_from(LogEntry)
                .where(LogEntry.operation_type == "ocr_export_success")
            )
        ).scalar_one()

        _job, artefact, po = await run_ocr_export(session=db_session, config=config)

        # Exactly one OCR_EXPORT_EVENT-PO landed.
        after_pos = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.OCR_EXPORT_EVENT.value)
            )
        ).scalar_one()
        assert after_pos == before_pos + 1
        assert str(po.po_type) == POType.OCR_EXPORT_EVENT.value
        # Atomicity: all canonical fields populated.
        for key in (
            "filename",
            "format",
            "sha256",
            "size_bytes",
            "ocr_revision_snapshot",
            "active_decision_event_uuids",
            "gate_mode",
            "active_stilprofil_version_uuid",
            "export_warnings",
            "export_attempt_id",
        ):
            assert key in po.payload, f"OCR_EXPORT_EVENT payload missing {key}"
        assert po.payload["sha256"] == artefact.sha256
        assert po.payload["export_attempt_id"] == attempt_id

        # Log entry on actually-started job.
        after_logs = (
            await db_session.execute(
                select(func.count())
                .select_from(LogEntry)
                .where(LogEntry.operation_type == "ocr_export_success")
            )
        ).scalar_one()
        assert after_logs == before_logs + 1

    async def test_failed_docx_writes_no_export_event_and_a_failed_log(
        self, db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """OCR-EXPORT_EVENT-Kein-Eintrag-Bei-Fehler-Test +
        Log-Eintrag-Bei-Gestarteten-Job-Test (failure side)."""
        project, _, _, _ = await _seed_project(db_session, n_segments=1)
        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=attempt_id,
        )
        await confirm_pflichtfragen(session=db_session, config=config)

        # Force build_ocr_docx to fail by monkeypatching the module-level
        # symbol that run_ocr_export imports.
        async def _broken_build(**kwargs):  # type: ignore[no-untyped-def]
            raise DocxArtefactFailed("synthetic failure")

        from waraq.ocr_export import service as ocr_export_service

        monkeypatch.setattr(ocr_export_service, "build_ocr_docx", _broken_build)

        before_pos = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.OCR_EXPORT_EVENT.value)
            )
        ).scalar_one()
        before_failed_logs = (
            await db_session.execute(
                select(func.count())
                .select_from(LogEntry)
                .where(LogEntry.operation_type == "ocr_export_failed")
            )
        ).scalar_one()

        with pytest.raises(DocxArtefactFailed):
            await run_ocr_export(session=db_session, config=config)

        after_pos = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.OCR_EXPORT_EVENT.value)
            )
        ).scalar_one()
        after_failed_logs = (
            await db_session.execute(
                select(func.count())
                .select_from(LogEntry)
                .where(LogEntry.operation_type == "ocr_export_failed")
            )
        ).scalar_one()
        # No OCR_EXPORT_EVENT, but a failed-log entry landed.
        assert after_pos == before_pos
        assert after_failed_logs == before_failed_logs + 1


# --- Snapshot rules ---------------------------------------------


@pytest.mark.asyncio
class TestSnapshotRules:
    async def test_snapshot_n_entries_matches_n_segments(self, db_session: AsyncSession) -> None:
        """OCR-Snapshot-Vollstaendigkeit-Test."""
        project, _, _, segments = await _seed_project(db_session, n_segments=3)
        # Force at least one segment to have a current_rev_uuid by writing
        # a Revision through the canonical service.
        from waraq.invariant.enums import OperationMode
        from waraq.revision import create_revision
        from waraq.schemas.enums import ChangeSource

        await create_revision(
            session=db_session,
            segment=segments[0],
            after_text="rev0",
            change_source=ChangeSource.OCR,
            operation_mode=OperationMode.AUTOMATIC,
        )

        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=attempt_id,
        )
        await confirm_pflichtfragen(session=db_session, config=config)
        _, _, po = await run_ocr_export(session=db_session, config=config)

        snapshot = po.payload["ocr_revision_snapshot"]
        assert len(snapshot) == 3  # one entry per exported segment

    async def test_active_de_uuids_excludes_glossary_and_preflight_and_old_attempts(
        self, db_session: AsyncSession
    ) -> None:
        """OCR-Decision-Snapshot-Allowlist-Test +
        OCR-Decision-Snapshot-Attempt-Bindung-Test."""
        project, _, _, _ = await _seed_project(db_session, n_segments=1)
        # 1) A glossary_management DE — must be EXCLUDED.
        de_glossary = await create_decision_event(
            session=db_session,
            scope_type=ScopeType.PROJECT,
            scope_uuid=project.project_uuid,
            decision_type="glossary_create",
            decision_source=DecisionSource.GLOSSARY_MANAGEMENT,
        )
        # 2) A style_management DE — must be EXCLUDED.
        de_style = await create_decision_event(
            session=db_session,
            scope_type=ScopeType.PROJECT,
            scope_uuid=project.project_uuid,
            decision_type="style_anchor_set",
            decision_source=DecisionSource.STYLE_MANAGEMENT,
        )
        # 3) An OLD export_confirmation DE (different attempt) — EXCLUDED.
        old_attempt_id = str(new_uuid())
        de_old_export = await create_decision_event(
            session=db_session,
            scope_type=ScopeType.PROJECT,
            scope_uuid=project.project_uuid,
            decision_type="ocr_export_pflichtfragen_bestaetigt",
            decision_source=DecisionSource.EXPORT_CONFIRMATION,
            related_export_attempt_id=old_attempt_id,
        )
        # 4) An ocr_review DE — INCLUDED.
        de_ocr = await create_decision_event(
            session=db_session,
            scope_type=ScopeType.PROJECT,
            scope_uuid=project.project_uuid,
            decision_type="ocr_review_no_go_to_go",
            decision_source=DecisionSource.OCR_REVIEW,
        )

        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),
            export_attempt_id=attempt_id,
        )
        # 5) The current attempt's confirmation DE — INCLUDED.
        de_current = await confirm_pflichtfragen(session=db_session, config=config)
        _, _, po = await run_ocr_export(session=db_session, config=config)

        active_set = set(po.payload["active_decision_event_uuids"])
        assert str(de_glossary.decision_event_uuid) not in active_set
        assert str(de_style.decision_event_uuid) not in active_set
        assert str(de_old_export.decision_event_uuid) not in active_set
        assert str(de_ocr.decision_event_uuid) in active_set
        assert str(de_current.decision_event_uuid) in active_set


# --- gate_mode reflected on the PO -----------------------------


@pytest.mark.asyncio
class TestGateModeOnPo:
    """OCR-Gate-Mode-Test."""

    async def test_go_with_warning_run_records_exportierbar_mit_warnungen(
        self, db_session: AsyncSession
    ) -> None:
        # Drive page to GO_WITH_WARNING via OCR review path.
        from waraq.ocr.review import (
            apply_findings_to_status,
            enter_in_review,
            make_default_severity_weights,
        )

        project, page, _, _ = await _seed_project(db_session, n_segments=1)
        weights = make_default_severity_weights()
        await record_ocr_error_instance(
            session=db_session,
            page_uuid=page.page_uuid,
            error_code=OcrErrorClass.F_03,  # hoch
        )
        await enter_in_review(session=db_session, page=page)
        await apply_findings_to_status(session=db_session, page=page, weights=weights)
        assert page.ocr_status == OcrStatus.GO_WITH_WARNING

        attempt_id = str(new_uuid())
        config = OcrExportConfig(
            project_uuid=project.project_uuid,
            pflichtfragen=_basic_pflichtfragen(),  # arbeitsstand
            export_attempt_id=attempt_id,
        )
        # In arbeitsstand mode, go_with_warning is exportable (with double-
        # confirmation flag set; the test exercises the run path which
        # does not require a separate confirmation DE in this v1.0
        # implementation — the Pflichtfragen DE doubles as the user's
        # "I accept" signal in arbeitsstand).
        await confirm_pflichtfragen(session=db_session, config=config)
        _, _, po = await run_ocr_export(session=db_session, config=config)
        assert po.payload["gate_mode"] == "exportierbar_mit_warnungen"


# --- Sprint-OCR §1.4 distinction: OCR_EXPORT_EVENT ≠ EXPORT_EVENT ---


class TestOcrExportEventDistinctFromExportEvent:
    """Sprint-OCR §1.4: OCR_EXPORT_EVENT and EXPORT_EVENT are different
    PO types with different semantics. They must never be silently
    mixed."""

    def test_two_distinct_canonical_potype_values(self) -> None:
        # Both exist in the enum.
        assert POType.EXPORT_EVENT.value == "export_event"
        assert POType.OCR_EXPORT_EVENT.value == "ocr_export_event"
        # And they are not equal.
        assert POType.EXPORT_EVENT.value != POType.OCR_EXPORT_EVENT.value


# --- Schema discipline ----------------------------------------


class TestOcrExportSchemaDiscipline:
    def test_decision_events_has_related_export_attempt_id_column(self) -> None:
        cols = DecisionEvent.__table__.columns
        assert "related_export_attempt_id" in cols
        assert cols["related_export_attempt_id"].nullable is True
