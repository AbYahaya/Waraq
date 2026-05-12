"""T-9.2.1 mandatory tests — Sprint 5 §4 (gate mode, format, logging,
state machine, user-action discipline).

Test ID coverage:
- Gate-Mode-Set-Correctly-Test
- Export-Warnings-Filled-When-With-Warnings-Test
- Pflichtfragen-Read-From-Decision-Events-Test
- Preflight-Recheck-At-Job-Start-Test
- Word-Kompatibel-Oeffnungs-Test
- RTL-Per-Run-Test
- Formatvorlagen-Baseline-Adherence-Test
- Log-Eintrag-Bei-Jedem-Versuch-Test
- Log-Eintrag-Vorabpruefung-Kein-Test
- Export-Starten-Decision-Event-Test
- Export-Starten-Nur-Aus-Exportierbar-Test
- Active-Decision-Event-Uuids-Is-Superseded-Filter-Test
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from tests.export._helpers import (
    accept_warnings,
    confirm_all_pflichtfragen,
    reach_exportierbar,
    seed_project_with_account,
    seed_segment_with_revision,
)
from waraq.audit.service import RuleFinding, record_befund
from waraq.audit.severity import default_severity_table
from waraq.export import (
    ExportConfig,
    ExportGateMode,
    ExportNotInExportableState,
    PreflightStateChanged,
    export_starten,
    run_export_job,
)
from waraq.identity import new_uuid
from waraq.preflight import (
    PreflightState,
    confirm_pflichtfrage,
    evaluate_preflight,
    save_export_profile_prefill,
    start_preflight_run,
)
from waraq.preflight.enums import WarningSlot
from waraq.schemas import (
    Job,
    LogEntry,
    ProvenanceObject,
)
from waraq.schemas.enums import DecisionSource, JobState, POType, ScopeType

# --- Gate-Mode-Set-Correctly-Test -------------------------------------


@pytest.mark.asyncio
class TestGateModeSetCorrectly:
    async def test_clean_project_gate_mode_exportierbar(self, db_session: AsyncSession) -> None:
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        run, state = await reach_exportierbar(db_session, project=project)
        assert state == PreflightState.EXPORTIERBAR

        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
        )
        result = await run_export_job(session=db_session, config=config)
        assert result.gate_mode == ExportGateMode.EXPORTIERBAR
        assert result.export_event_po.payload["gate_mode"] == "exportierbar"

    async def test_warnings_project_gate_mode_with_warnings(self, db_session: AsyncSession) -> None:
        project, account_uuid = await seed_project_with_account(db_session)
        seg = await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        # Inject an open mittel Befund → W-01 warning.
        audit_job = Job(
            job_uuid=new_uuid(),
            project_uuid=project.project_uuid,
            job_type="audit",
            state=JobState.COMPLETED.value,
        )
        db_session.add(audit_job)
        await db_session.flush()
        await record_befund(
            session=db_session,
            finding=RuleFinding(
                regelkennung="D-01",
                satz_uuid=seg.satz_uuid,
                detection_context={},
            ),
            project_uuid=project.project_uuid,
            audit_run_job_uuid=audit_job.job_uuid,
            severity_table=default_severity_table(),
        )

        run = await start_preflight_run(session=db_session, project_uuid=project.project_uuid)
        await confirm_all_pflichtfragen(
            db_session, project_uuid=project.project_uuid, preflight_run=run
        )
        # Accept the W-01 warning.
        await accept_warnings(
            db_session,
            project_uuid=project.project_uuid,
            preflight_run=run,
            warning_slots=[WarningSlot.W_01_MITTEL_AUDIT],
        )
        ev = await evaluate_preflight(
            session=db_session,
            project_uuid=project.project_uuid,
            preflight_run=run,
        )
        assert ev.state == PreflightState.EXPORTIERBAR_MIT_WARNUNGEN

        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
            export_warnings=[{"slot": "w_01_mittel_audit", "befund_uuid": "x"}],
        )
        result = await run_export_job(session=db_session, config=config)
        assert result.gate_mode == ExportGateMode.EXPORTIERBAR_MIT_WARNUNGEN
        assert result.export_event_po.payload["gate_mode"] == "exportierbar_mit_warnungen"
        assert len(result.export_event_po.payload["export_warnings"]) >= 1


# --- Export-Warnings-Filled-When-With-Warnings-Test -------------------


@pytest.mark.asyncio
class TestExportWarningsFilled:
    async def test_export_warnings_empty_when_exportierbar(self, db_session: AsyncSession) -> None:
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        run, _state = await reach_exportierbar(db_session, project=project)
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
            export_warnings=[{"would_be": "ignored"}],  # caller passes nothing-valid
        )
        result = await run_export_job(session=db_session, config=config)
        # Gate is `exportierbar` → warnings array MUST be empty regardless of
        # what the caller passed (the service forces `[]` when no warnings).
        assert result.export_event_po.payload["export_warnings"] == []


# --- Pflichtfragen-Read-From-Decision-Events-Test ---------------------


@pytest.mark.asyncio
class TestPflichtfragenFromDecisionEvents:
    async def test_export_config_built_from_des_not_from_profile(
        self, db_session: AsyncSession
    ) -> None:
        """Saved profile pre-fills with 'PROFILE' answers; active
        confirmations record 'ACTIVE'. The export_config payload MUST
        contain the ACTIVE answers — never the profile answers."""
        from tests.preflight._helpers import canonical_pflichtfrage_payload

        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")

        # Helper: the PROFILE pre-fill carries a deliberately *different*
        # but still canonical-shape answer than the ACTIVE confirmation,
        # so the test can verify the export reads the ACTIVE one.
        def _profile_payload(idx: int) -> tuple[str, dict[str, object]]:
            key, _ = canonical_pflichtfrage_payload(idx)
            if idx in (1, 2):
                return key, {"heading_level": 6}
            if idx == 3:
                return key, {"position": "back"}
            return key, {"display": False}

        # Save profile pre-fills with the "stale" canonical-shape answers.
        for i in range(1, 5):
            key, ans = _profile_payload(i)
            await save_export_profile_prefill(
                session=db_session,
                project_uuid=project.project_uuid,
                frage_index=i,
                frage_key=key,
                prefilled_answer=ans,
            )
        run = await start_preflight_run(session=db_session, project_uuid=project.project_uuid)
        # Active confirmations with the canonical default (different from profile).
        for i in range(1, 5):
            key, ans = canonical_pflichtfrage_payload(i)
            await confirm_pflichtfrage(
                session=db_session,
                project_uuid=project.project_uuid,
                preflight_run_uuid=run.job_uuid,
                frage_index=i,
                frage_key=key,
                answer=ans,
            )

        # The Sprint-4 evaluator counts confirmations against the
        # `preflight_run_uuid`. The export-side pflichtfragen lookup
        # uses the export_attempt_id. Confirm pflichtfragen against the
        # export_attempt_id explicitly to drive the test through the
        # code path Pflichtfragen-Read-From-Decision-Events-Test
        # exercises (export-side query).
        export_attempt_id = str(new_uuid())
        for i in range(1, 5):
            key, ans = canonical_pflichtfrage_payload(i)
            await confirm_pflichtfrage(
                session=db_session,
                project_uuid=project.project_uuid,
                preflight_run_uuid=run.job_uuid,
                frage_index=i,
                frage_key=key,
                answer=ans,
            )

        # Re-write a confirmation tagged with the EXPORT attempt id.
        from waraq.decisions import create_decision_event

        for i in range(1, 5):
            key, ans = canonical_pflichtfrage_payload(i)
            await create_decision_event(
                session=db_session,
                scope_type=ScopeType.PROJECT,
                scope_uuid=project.project_uuid,
                decision_type="pflichtfrage_bestaetigung",
                decision_source=DecisionSource.PREFLIGHT_CONFIRMATION,
                content={
                    "frage_index": i,
                    "frage_key": key,
                    "answer": ans,
                },
                related_export_attempt_id=export_attempt_id,
            )

        ev = await evaluate_preflight(
            session=db_session,
            project_uuid=project.project_uuid,
            preflight_run=run,
        )
        assert ev.state == PreflightState.EXPORTIERBAR
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=export_attempt_id,
            preflight_run=run,
        )
        result = await run_export_job(session=db_session, config=config)
        export_config = result.export_event_po.payload["export_config"]
        pf_list = export_config["pflichtfragen"]
        assert len(pf_list) == 4
        for entry in pf_list:
            idx = entry["frage_index"]
            _, expected_active = canonical_pflichtfrage_payload(idx)
            _, profile = _profile_payload(idx)
            assert entry["answer"] == expected_active
            # Critically — never the profile-pre-fill answer.
            assert entry["answer"] != profile


# --- Preflight-Recheck-At-Job-Start-Test ------------------------------


@pytest.mark.asyncio
class TestPreflightRecheckAtJobStart:
    async def test_state_change_between_user_action_and_job_fails_job(
        self, db_session: AsyncSession
    ) -> None:
        project, account_uuid = await seed_project_with_account(db_session)
        seg = await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        run, state = await reach_exportierbar(db_session, project=project)
        assert state == PreflightState.EXPORTIERBAR

        # User issues export_starten — passes at the entry check.
        _de, attempt_id = await export_starten(
            session=db_session,
            project_uuid=project.project_uuid,
            preflight_state=state,
        )

        # Between user action and job start, a critical audit finding appears.
        audit_job = Job(
            job_uuid=new_uuid(),
            project_uuid=project.project_uuid,
            job_type="audit",
            state=JobState.COMPLETED.value,
        )
        db_session.add(audit_job)
        await db_session.flush()
        await record_befund(
            session=db_session,
            finding=RuleFinding(
                regelkennung="C-01",  # kritisch → P-03 → blockiert
                satz_uuid=seg.satz_uuid,
                detection_context={},
            ),
            project_uuid=project.project_uuid,
            audit_run_job_uuid=audit_job.job_uuid,
            severity_table=default_severity_table(),
        )

        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=attempt_id,
            preflight_run=run,
        )
        # Snapshot existing EXPORT_EVENT count — interactive dev sessions
        # may leave POs around from real exports against the same DB.
        # The recheck contract is "no NEW EXPORT_EVENT on state change",
        # not "global table is empty".
        po_count_before = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.EXPORT_EVENT.value)
            )
        ).scalar_one()
        with pytest.raises(PreflightStateChanged):
            await run_export_job(session=db_session, config=config)

        # No NEW EXPORT_EVENT row.
        po_count_after = (
            await db_session.execute(
                select(func.count())
                .select_from(ProvenanceObject)
                .where(ProvenanceObject.po_type == POType.EXPORT_EVENT.value)
            )
        ).scalar_one()
        assert po_count_after == po_count_before
        # FAILED Log-Eintrag with reason `preflight_state_changed`.
        log_rows = (
            (
                await db_session.execute(
                    select(LogEntry)
                    .where(LogEntry.operation_type == "export_failed")
                    .where(LogEntry.scope_uuid == project.project_uuid)
                )
            )
            .scalars()
            .all()
        )
        assert len(log_rows) == 1
        assert log_rows[0].result.get("reason") == "preflight_state_changed"


# --- Export-Starten tests ---------------------------------------------


@pytest.mark.asyncio
class TestExportStartenDecisionEvent:
    async def test_export_starten_writes_distinct_decision_event(
        self, db_session: AsyncSession
    ) -> None:
        project, _ = await seed_project_with_account(db_session)
        de, attempt_id = await export_starten(
            session=db_session,
            project_uuid=project.project_uuid,
            preflight_state=PreflightState.EXPORTIERBAR,
        )
        assert de.decision_type == "exportstart"
        assert de.scope_type == ScopeType.PROJECT.value
        assert de.scope_uuid == project.project_uuid
        # Has its own export_attempt_id used as related_export_attempt_id.
        assert de.related_export_attempt_id == attempt_id
        # Distinct from per-warning confirmations: those use
        # decision_type='preflight_warning_accepted_*'; this one is
        # 'exportstart'.
        assert "preflight_warning_accepted_" not in de.decision_type


@pytest.mark.asyncio
class TestExportStartenOnlyFromExportable:
    async def test_export_starten_refused_from_blockiert(self, db_session: AsyncSession) -> None:
        project, _ = await seed_project_with_account(db_session)
        # Refuse before any DB writes — no Job, no Log-Eintrag.
        log_count_before = (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one()
        with pytest.raises(ExportNotInExportableState):
            await export_starten(
                session=db_session,
                project_uuid=project.project_uuid,
                preflight_state=PreflightState.BLOCKIERT,
            )
        log_count_after = (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one()
        assert log_count_after == log_count_before


# --- Log-Eintrag tests ------------------------------------------------


@pytest.mark.asyncio
class TestLogEintragOnEveryAttempt:
    async def test_success_attempt_logs_export_success(self, db_session: AsyncSession) -> None:
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        run, _ = await reach_exportierbar(db_session, project=project)
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
        )
        await run_export_job(session=db_session, config=config)
        rows = (
            (
                await db_session.execute(
                    select(LogEntry)
                    .where(LogEntry.operation_type == "export_success")
                    .where(LogEntry.scope_uuid == project.project_uuid)
                )
            )
            .scalars()
            .all()
        )
        assert len(rows) == 1


class TestLogEintragVorabpruefungKein:
    """Code-review test: `export_starten` raises BEFORE any Log-Eintrag
    is written. Pre-checks produce no log noise."""

    def test_export_starten_does_not_call_log_event_in_refusal_path(self) -> None:
        import inspect

        from waraq.export import service as svc

        src = inspect.getsource(svc.export_starten)
        # `log_event` MUST NOT be called inside `export_starten` (the
        # function's only DE write goes through `create_decision_event`).
        assert "log_event" not in src, (
            "export_starten calls log_event; pre-check refusal must not "
            "produce a Log-Eintrag (Sprint 5 R-S5-12)."
        )


# --- Word-compatibility tests -----------------------------------------


@pytest.mark.asyncio
class TestWordCompatibility:
    async def test_artefact_is_valid_docx(self, db_session: AsyncSession) -> None:
        """Smoke test: artefact bytes round-trip through python-docx
        without errors. Word-compatible per OpenXML reading."""
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(
            db_session,
            project=project,
            text="إن الحمد لله\n---\nLob sei Gott",
        )
        run, _ = await reach_exportierbar(db_session, project=project)
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
        )
        # Use a custom store that retains bytes for the round-trip check.
        from waraq.export import InMemoryArtefactStore

        store = InMemoryArtefactStore()
        result = await run_export_job(session=db_session, config=config, artefact_store=store)
        bytes_ = store.get(artefact_uuid=result.artefact_uuid)
        assert bytes_ is not None
        # python-docx must be able to re-open without raising.
        doc = Document(io.BytesIO(bytes_))
        # At least one paragraph (project title heading) was written.
        assert len(doc.paragraphs) >= 1


@pytest.mark.asyncio
class TestRtlPerRun:
    async def test_arabic_paragraph_carries_bidi(self, db_session: AsyncSession) -> None:
        """Per Formatvorlagen-Baseline v1.1 §7.2: Arabic paragraphs
        carry `<w:bidi/>` at the paragraph-properties level. Sprint 5
        RTL-Per-Run-Test asserts this structurally."""
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(
            db_session,
            project=project,
            text="إن الحمد لله\n---\nLob sei Gott",
        )
        run, _ = await reach_exportierbar(db_session, project=project)
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
        )
        from waraq.export import InMemoryArtefactStore

        store = InMemoryArtefactStore()
        result = await run_export_job(session=db_session, config=config, artefact_store=store)
        bytes_ = store.get(artefact_uuid=result.artefact_uuid)
        doc = Document(io.BytesIO(bytes_))
        # Find at least one paragraph with the AR text and confirm bidi mark.
        ar_para = None
        for p in doc.paragraphs:
            if "إن الحمد" in p.text:
                ar_para = p
                break
        assert ar_para is not None
        # The bidi element is on the paragraph-properties level.
        from docx.oxml.ns import qn

        p_pr = ar_para._p.find(qn("w:pPr"))
        assert p_pr is not None
        bidi = p_pr.find(qn("w:bidi"))
        assert bidi is not None, "Arabic paragraph missing <w:bidi/>"


@pytest.mark.asyncio
class TestFormatvorlagenAdherence:
    async def test_toc_field_present(self, db_session: AsyncSession) -> None:
        """Per Formatvorlagen-Baseline v1.1 §7.2 — TOC `\\o "1-6"`
        (Schluss-Audit Paket 7 Item 2 (a), 2026-05-08)."""
        project, account_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=project, text="x\n---\ny")
        run, _ = await reach_exportierbar(db_session, project=project)
        config = ExportConfig(
            project_uuid=project.project_uuid,
            account_uuid=account_uuid,
            project_title="T",
            current_export_attempt_id=str(new_uuid()),
            preflight_run=run,
        )
        from waraq.export import InMemoryArtefactStore

        store = InMemoryArtefactStore()
        result = await run_export_job(session=db_session, config=config, artefact_store=store)
        bytes_ = store.get(artefact_uuid=result.artefact_uuid)
        # The TOC field instruction text must be present.
        # python-docx serializes field instructions in the document XML.
        from docx import Document

        doc = Document(io.BytesIO(bytes_))
        body_xml = doc.element.body.xml  # type: ignore[attr-defined]
        assert 'TOC \\o "1-6"' in body_xml
