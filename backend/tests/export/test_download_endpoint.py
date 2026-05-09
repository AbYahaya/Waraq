"""M5 — Translation-export download endpoint tests.

Verifies:
- GET /exports/artefacts/{po_uuid} streams a valid DOCX rebuilt from
  the EXPORT_EVENT-PO's `revision_snapshot[]`.
- Snapshot fidelity: the rebuilt DOCX captures `Revision.after_text`
  (immutable per H-5), not current `Segment.text_content`.
- 404 for non-EXPORT_EVENT po_uuids and for cross-account access.
- Read-only: download writes nothing.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from tests.conftest import seed_account_uuid
from tests.export._helpers import (
    seed_project_with_account,
    seed_segment_with_revision,
)
from waraq.api.dependencies import get_db_session
from waraq.api.main import create_app
from waraq.auth.tokens import issue_token
from waraq.export import ExportConfig, run_export_job
from waraq.identity import new_uuid
from waraq.invariant.enums import OperationMode
from waraq.preflight import (
    PFLICHTFRAGE_COUNT,
    confirm_pflichtfrage,
    evaluate_preflight,
    start_preflight_run,
)
from waraq.revision import create_revision
from waraq.schemas import (
    DecisionEvent,
    LogEntry,
    ProvenanceObject,
    Revision,
)
from waraq.schemas.enums import ChangeSource


@pytest.fixture
async def authed_client(db_session: AsyncSession):
    account_uuid = new_uuid()
    await seed_account_uuid(db_session, account_uuid)
    app = create_app()

    async def _override():
        yield db_session

    app.dependency_overrides[get_db_session] = _override
    token = issue_token(account_uuid=account_uuid)
    headers = {"Authorization": f"Bearer {token}"}
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        yield client, account_uuid, headers


async def _seed_project_for_account(
    db_session: AsyncSession, account_uuid
) -> tuple[object, object]:
    """Seed a project owned by the authed account + one segment with a Revision."""
    from waraq.schemas import Project

    project = Project(project_uuid=new_uuid(), account_uuid=account_uuid, name="t")
    db_session.add(project)
    await db_session.flush()
    seg = await seed_segment_with_revision(
        db_session,
        project=project,
        text="إن الحمد لله\n---\nLob sei Gott",
    )
    return project, seg


async def _run_export(db_session: AsyncSession, project, account_uuid) -> object:
    run = await start_preflight_run(session=db_session, project_uuid=project.project_uuid)
    for i in range(1, PFLICHTFRAGE_COUNT + 1):
        await confirm_pflichtfrage(
            session=db_session,
            project_uuid=project.project_uuid,
            preflight_run_uuid=run.job_uuid,
            frage_index=i,
            frage_key=f"frage_{i}",
            answer={"value": "yes"},
        )
    await evaluate_preflight(
        session=db_session, project_uuid=project.project_uuid, preflight_run=run
    )
    config = ExportConfig(
        project_uuid=project.project_uuid,
        account_uuid=account_uuid,
        project_title="Test Project",
        current_export_attempt_id=str(new_uuid()),
        preflight_run=run,
    )
    return await run_export_job(session=db_session, config=config)


@pytest.mark.asyncio
class TestDownloadEndpoint:
    async def test_download_returns_valid_docx(
        self, db_session: AsyncSession, authed_client
    ) -> None:
        client, account_uuid, headers = authed_client
        project, _seg = await _seed_project_for_account(db_session, account_uuid)
        result = await _run_export(db_session, project, account_uuid)

        resp = await client.get(
            f"/exports/artefacts/{result.export_event_po.po_uuid}", headers=headers
        )
        assert resp.status_code == 200, resp.text
        assert resp.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument"
        )
        assert "attachment" in resp.headers.get("content-disposition", "")

        # DOCX bytes round-trip through python-docx.
        doc = Document(io.BytesIO(resp.content))
        # Title heading present.
        assert any(p.text == "Test Project" for p in doc.paragraphs)
        # Arabic text from the export present.
        assert any("إن الحمد" in p.text for p in doc.paragraphs)

    async def test_download_404_for_unknown_po(
        self, db_session: AsyncSession, authed_client
    ) -> None:
        client, _account_uuid, headers = authed_client
        resp = await client.get(f"/exports/artefacts/{new_uuid()}", headers=headers)
        assert resp.status_code == 404

    async def test_download_404_for_non_export_event_po(
        self, db_session: AsyncSession, authed_client
    ) -> None:
        """OCR_EXPORT_EVENT POs use a different download endpoint —
        the translation download must refuse them."""
        from tests.readout._helpers import seed_po
        from waraq.schemas.enums import POType, ScopeType

        client, account_uuid, headers = authed_client
        project, _seg = await _seed_project_for_account(db_session, account_uuid)
        ocr_po = await seed_po(
            db_session,
            po_type=POType.OCR_EXPORT_EVENT,
            scope_type=ScopeType.PROJECT,
            scope_uuid=project.project_uuid,
        )
        resp = await client.get(f"/exports/artefacts/{ocr_po.po_uuid}", headers=headers)
        assert resp.status_code == 404

    async def test_download_404_for_other_accounts_export(
        self, db_session: AsyncSession, authed_client
    ) -> None:
        client, _account_uuid, headers = authed_client
        # Seed a different account's project + export.
        other_account, other_uuid = await seed_project_with_account(db_session)
        await seed_segment_with_revision(db_session, project=other_account, text="x\n---\ny")
        result = await _run_export(db_session, other_account, other_uuid)

        # Authed user is `account_uuid`, not `other_uuid` — should 404.
        resp = await client.get(
            f"/exports/artefacts/{result.export_event_po.po_uuid}", headers=headers
        )
        assert resp.status_code == 404


@pytest.mark.asyncio
class TestSnapshotFidelity:
    async def test_rebuilt_docx_uses_frozen_revision_text(
        self, db_session: AsyncSession, authed_client
    ) -> None:
        """After export, mutate the segment to a NEW revision. The
        download MUST still show the OLD text frozen in
        `revision_snapshot[]` — H-5 immutability is the load-bearing
        invariant here."""
        client, account_uuid, headers = authed_client
        project, seg = await _seed_project_for_account(db_session, account_uuid)
        result = await _run_export(db_session, project, account_uuid)

        # Mutate segment to v2 — new Revision row, segment.text_content updated.
        await create_revision(
            session=db_session,
            segment=seg,
            after_text="جديد\n---\nNeu",
            change_source=ChangeSource.MANUAL,
            operation_mode=OperationMode.MANUAL_WITH_CONFIRMATION,
        )
        await db_session.refresh(seg)
        assert "Neu" in (seg.text_content or "")

        resp = await client.get(
            f"/exports/artefacts/{result.export_event_po.po_uuid}", headers=headers
        )
        assert resp.status_code == 200
        doc = Document(io.BytesIO(resp.content))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Original (frozen) text present.
        assert "Lob sei Gott" in all_text
        # New (post-export) text MUST NOT leak in.
        assert "Neu" not in all_text


@pytest.mark.asyncio
class TestDownloadIsReadOnly:
    async def test_download_writes_nothing(self, db_session: AsyncSession, authed_client) -> None:
        client, account_uuid, headers = authed_client
        project, _seg = await _seed_project_for_account(db_session, account_uuid)
        result = await _run_export(db_session, project, account_uuid)

        rev_count_before = (
            await db_session.execute(select(func.count()).select_from(Revision))
        ).scalar_one()
        de_count_before = (
            await db_session.execute(select(func.count()).select_from(DecisionEvent))
        ).scalar_one()
        log_count_before = (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one()
        po_count_before = (
            await db_session.execute(select(func.count()).select_from(ProvenanceObject))
        ).scalar_one()

        await client.get(f"/exports/artefacts/{result.export_event_po.po_uuid}", headers=headers)

        assert (
            await db_session.execute(select(func.count()).select_from(Revision))
        ).scalar_one() == rev_count_before
        assert (
            await db_session.execute(select(func.count()).select_from(DecisionEvent))
        ).scalar_one() == de_count_before
        assert (
            await db_session.execute(select(func.count()).select_from(LogEntry))
        ).scalar_one() == log_count_before
        assert (
            await db_session.execute(select(func.count()).select_from(ProvenanceObject))
        ).scalar_one() == po_count_before
